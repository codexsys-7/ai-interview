# This is my final version of MVP

from dotenv import load_dotenv
load_dotenv()

from ats import compute_ats_score


import io, json, os, re, logging, tempfile, time, random
from typing import Any, Dict, List, Optional
from uuid import UUID
from contextlib import asynccontextmanager

import fitz          
import docx
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Response, Depends
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, EmailStr
from openai import OpenAI
from openai import RateLimitError

from typing import Optional

from dotenv import load_dotenv
from pathlib import Path


from sqlmodel import Session, select
from db import engine, init_db, InterviewSession, InterviewAnswer, JobDescription
from services.embedding_service import generate_embedding, find_similar_answers
from services.interview_orchestrator import InterviewOrchestrator
from services.conversation_context import (
    get_all_answers,
    build_conversation_summary,
    detect_repeated_topics,
    extract_topics
)
from services.contradiction_detector import detect_contradictions
from services.tts_service import (
    TTSService,
    get_tts_service,
    generate_interview_speech
)
from services.job_introduction_generator import (
    JobIntroductionGenerator,
    get_job_introduction_generator,
    generate_job_introduction,
    IntroductionMode
)
from models import Resume, User

from passlib.context import CryptContext
from jose import jwt
from datetime import datetime, timedelta

from fastapi.middleware.cors import CORSMiddleware



# Load .env from this file's directory and OVERRIDE any existing env vars
env_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=env_path, override=True)


# ---------- Lifespan Event Handler ----------
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Create tables if they do not exist
    init_db()
    yield
    # Shutdown: Cleanup code would go here if needed


# ---------- FastAPI & CORS ----------
app = FastAPI(lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "null",  # Allow file:// protocol for testing
    ],
    allow_origin_regex=r"http://(localhost|127\.0\.0\.1):(5173|8000)",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)



@app.options("/{path:path}")
def preflight_handler(path: str):
    return Response(status_code=200)


@app.options("/{path:path}")
def options_handler(path: str):
    return Response(status_code=200)


@app.get("/api/__debug")
def debug():
    return {
        "ok": True,
        "file": __file__,
        "origins": ["http://localhost:5173", "http://127.0.0.1:5173"],
    }



# These are my Config + helpers
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

JWT_SECRET = os.getenv("JWT_SECRET", "dev_secret_change_me")
JWT_ALG = "HS256"
JWT_EXPIRE_MIN = 60 * 24 * 7  # 7 days

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(password: str, password_hash: str) -> bool:
    return pwd_context.verify(password, password_hash)

def create_access_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.utcnow() + timedelta(minutes=JWT_EXPIRE_MIN),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


class SignupReq(BaseModel):
    full_name: str
    email: EmailStr
    password: str

class LoginReq(BaseModel):
    email: EmailStr
    password: str

class AuthResp(BaseModel):
    token: str
    user: dict


# These are my Longin and Sign Up page Endpoints.
@app.post("/api/auth/signup", response_model=AuthResp)
def signup(req: SignupReq):
    if len(req.password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters.")

    # bcrypt limit: 72 bytes (not chars)
    if len(req.password.encode("utf-8")) > 100:
        raise HTTPException(status_code=400, detail="Password too long (max 100 bytes). Use a shorter password.")


    with Session(engine) as session:
        existing = session.exec(select(User).where(User.email == req.email.lower())).first()
        if existing:
            raise HTTPException(status_code=409, detail="Email already registered.")

        user = User(
            full_name=req.full_name.strip(),
            email=req.email.lower().strip(),
            password_hash=hash_password(req.password),
        )
        session.add(user)
        session.commit()
        session.refresh(user)

        token = create_access_token(user.id, user.email)
        return AuthResp(token=token, user={"id": user.id, "full_name": user.full_name, "email": user.email})


@app.post("/api/auth/login", response_model=AuthResp)
def login(req: LoginReq):
    with Session(engine) as session:
        user = session.exec(select(User).where(User.email == req.email.lower())).first()

        if not user or not verify_password(req.password, user.password_hash):
            raise HTTPException(status_code=401, detail="Invalid email or password.")

        if not user.is_active:
            raise HTTPException(status_code=403, detail="Account disabled.")

        token = create_access_token(user.id, user.email)
        return AuthResp(token=token, user={"id": user.id, "full_name": user.full_name, "email": user.email})


# ---------- OpenAI client ----------
OPENAI_KEY = os.environ.get("OPENAI_API_KEY")
print("OPENAI KEY DEBUG:", (OPENAI_KEY or "MISSING")[:8], "...", (OPENAI_KEY or "")[-4:])

client = OpenAI(api_key=OPENAI_KEY)


def call_openai_with_backoff(*, messages, model="gpt-4o-mini", response_format=None, temperature=0.3):
    """Retries OpenAI calls with exponential backoff to avoid 429 rate-limit errors."""
    for attempt in range(5):
        try:
            return client.chat.completions.create(
                model=model,
                temperature=temperature,
                response_format=response_format,
                messages=messages,
            )
        except RateLimitError as e:
            wait = min(20, 2 ** attempt) + random.random()
            print(f"[WARN] Rate limit hit. Retrying in {wait:.1f}s...")
            time.sleep(wait)
    # Final attempt (let it raise if still failing)
    return client.chat.completions.create(
        model=model,
        temperature=temperature,
        response_format=response_format,
        messages=messages,
    )


SYSTEM_PROMPT = """
You are an ATS + resume analysis engine.

Your job is to ANALYZE THE GIVEN RESUME TEXT ONLY and return STRICT JSON.

Rules:
- Use ONLY information that is explicitly present or clearly implied in the resume.
- Avoid generic guesses like "team player", "good communication", etc. unless those words exist.
- Pay close attention to:
  - Tools, libraries, frameworks, languages, platforms
  - Domains (finance, banking, healthcare, e-commerce, technology, telecommunications, advertising, media and entertainment, manufacturing and industrial, textile and fashion, energy, oil and gas, transportation, government and public sector, education, real estate and construction, hospitality, gaming, robotics, cybersecurity.)
  - Seniority level (intern, junior, senior, lead, architect)
  - Measurable impact (%, $, time saved) if present
- If something is unclear, LEAVE IT OUT instead of hallucinating.

Schema expectations:
- skills:
  - 5-30 items.
  - Each item must be a specific skill or tool actually supported by the text.
  - Examples: "Python", "FastAPI", "Kafka", "Time-series forecasting", "ATS optimization".
- fallbackRoles:
  - Up to 4 realistic role titles based on the experience in the resume.
  - Examples: "Data Scientist", "ML Engineer", "Python Developer", "Business Analyst".
- rare:
  - readability, applicability, remarkability, total — floats between 0 and 5 with 1 decimal.
  - total must be the average of the first three, rounded to 1 decimal.
  - Be discriminative: 3.0 is "okay", 4.0 is "strong", 4.8+ is "exceptional".
- atsScore:
  - 0-100.
  - Base it on:
    - Keyword density vs skills/experience,
    - Clarity of section headings,
    - Simplicity of layout (no tables/columns),
    - Match between skills and likely target roles.
- atsSuggestions:
  - If atsScore < 90: return 3-7 concrete, resume-specific suggestions.
  - Each suggestion must mention something from THIS resume (a missing keyword, weak bullet, etc.).
  - If atsScore ≥ 90: return [].
- keywords:
  - 3-15 short phrases.
  - Use domain + tech combos where possible, e.g. "healthcare analytics", "ETL pipelines", "LLM-powered chatbots".

Response:
- Return ONLY compact JSON that conforms exactly to the provided schema.
- No markdown, no prose outside JSON.
""".strip()


SYSTEM_PROMPT_JD = """
You are an ATS + job-match engine.

You will receive:
- RESUME TEXT
- JOB DESCRIPTION TEXT

You MUST compare the resume against the job description and return STRICT JSON ONLY.

Rules:
- Use ONLY info that exists in the resume or JD. No hallucinations.
- skills: 5-30 items from the resume ONLY (tools, frameworks, methods, domain skills).
- keywords: 5-20 important phrases derived from the JD + resume.
- fallbackRoles: up to 4 roles based primarily on the resume.
- rare: 0-5 with 1 decimal (readability, applicability, remarkability, total=avg rounded 1 decimal).
  * "applicability" MUST reflect match to the job description when JD is present.
- atsScore: 0-100. This MUST reflect how well the resume matches the job description.

ATS rubric (match-based):
- Searchability (20): clean headings, contact signals, standard sections inferred from text.
- Hard skills match (35): tools/tech overlap with JD requirements.
- Responsibilities match (25): resume evidence for JD responsibilities.
- Seniority match (10): job level alignment.
- Recruiter tips (10): metrics, impact, clarity, web presence.

Also include:
- jobMatchScore: 0-100 (can be same as atsScore, but include it)
- matchedKeywords: 5-20 strings found in both resume and JD
- missingKeywords: 5-30 important JD terms not found in resume

atsSuggestions:
- If atsScore < 90 return 4-8 JD-specific suggestions.
- Suggestions MUST reference missing keywords/responsibilities from this JD.
- If atsScore >= 90 return [].

- rare MUST be a JSON object with keys: readability, applicability, remarkability, total.
  Example: "rare": {"readability": 4.2, "applicability": 4.0, "remarkability": 3.8, "total": 4.0}
  NEVER return rare as a number.


Return ONLY JSON. No markdown. No extra keys outside schema. no prose outside JSON.
""".strip()



# ---------- Helpers ----------
def extract_pdf_text(data: bytes) -> str:
    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(p.get_text() for p in doc)

def extract_docx_text(data: bytes) -> str:
    d = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in d.paragraphs)

def clean_text(t: str) -> str:
    t = t.replace("\x00", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()

# ------------------- Question generation helpers (LLM) -------------------

# Comprehensive interviewer profiles with realistic question focus areas
# Each interviewer has specific areas they focus on in real interviews
INTERVIEWER_PROFILES = {
    "HR": {
        "title": "Human Resources",
        "focus_areas": [
            "culture fit and values alignment",
            "teamwork and collaboration style",
            "conflict resolution and interpersonal skills",
            "career goals and motivation",
            "communication skills",
            "work-life balance expectations",
            "salary expectations and benefits",
            "reason for leaving previous role",
            "handling feedback and criticism",
            "diversity and inclusion awareness"
        ],
        "question_types": {"behavioral": 60, "situational": 25, "cultural": 15},
        "tone": "warm, conversational, assessing personality and culture fit",
        "example_questions": [
            "Tell me about a time you had a disagreement with a colleague. How did you handle it?",
            "What kind of work environment helps you do your best work?",
            "How do you handle constructive criticism?",
            "Where do you see yourself in 3-5 years?",
            "Why are you interested in leaving your current position?"
        ]
    },
    "Manager": {
        "title": "Hiring Manager",
        "focus_areas": [
            "day-to-day work execution and delivery",
            "prioritization and time management",
            "collaboration with team members",
            "handling deadlines and pressure",
            "communication with stakeholders",
            "problem-solving approach",
            "ownership and accountability",
            "project management and planning",
            "technical skills relevant to the role",
            "mentoring and helping others"
        ],
        "question_types": {"behavioral": 40, "technical": 35, "situational": 25},
        "tone": "practical, focused on real work scenarios and team dynamics",
        "example_questions": [
            "Walk me through how you would approach a project with a tight deadline.",
            "Tell me about a time you had to deliver something with incomplete requirements.",
            "How do you prioritize when you have multiple urgent tasks?",
            "Describe a situation where you had to push back on a stakeholder request.",
            "How do you keep your team informed about project progress?"
        ]
    },
    "Tech Lead": {
        "title": "Technical Lead / Senior Engineer",
        "focus_areas": [
            "deep technical knowledge and expertise",
            "system design and architecture",
            "code quality and best practices",
            "debugging and troubleshooting complex issues",
            "performance optimization",
            "security considerations",
            "technical trade-offs and decision making",
            "staying current with technology trends",
            "code review and mentoring",
            "technical documentation"
        ],
        "question_types": {"technical": 75, "problem_solving": 15, "behavioral": 10},
        "tone": "technical, detailed, probing for depth of knowledge",
        "example_questions": [
            "How would you design a system that handles 10 million requests per day?",
            "Explain a complex technical problem you solved and your approach.",
            "What's your process for debugging a production issue at 2 AM?",
            "How do you decide between building vs buying a solution?",
            "Walk me through your code review process."
        ]
    },
    "CEO": {
        "title": "Chief Executive Officer",
        "focus_areas": [
            "business impact and value creation",
            "strategic thinking and vision",
            "leadership and influence",
            "understanding of company mission",
            "innovation and creative thinking",
            "risk assessment and management",
            "growth mindset and adaptability",
            "customer-centric thinking",
            "handling ambiguity and change",
            "long-term thinking"
        ],
        "question_types": {"strategic": 50, "behavioral": 30, "vision": 20},
        "tone": "big-picture, strategic, assessing leadership potential and business acumen",
        "example_questions": [
            "How would your work contribute to our company's mission?",
            "Tell me about a time you identified a significant business opportunity.",
            "How do you stay ahead of industry trends?",
            "What's the biggest risk you've taken professionally, and what was the outcome?",
            "How do you make decisions when you don't have all the information?"
        ]
    },
    "CFO": {
        "title": "Chief Financial Officer",
        "focus_areas": [
            "cost-benefit analysis and ROI thinking",
            "budget management and resource allocation",
            "financial impact of decisions",
            "efficiency and optimization",
            "data-driven decision making",
            "risk vs reward assessment",
            "understanding of business metrics",
            "vendor management and negotiations",
            "compliance and governance",
            "scaling and growth considerations"
        ],
        "question_types": {"analytical": 50, "behavioral": 30, "business": 20},
        "tone": "analytical, numbers-focused, assessing business and financial acumen",
        "example_questions": [
            "How do you measure the success of your projects in terms of business value?",
            "Tell me about a time you had to make a decision with budget constraints.",
            "How do you prioritize investments when resources are limited?",
            "Describe a situation where you identified cost savings.",
            "How do you calculate the ROI of a technical initiative?"
        ]
    },
    "President": {
        "title": "President / COO",
        "focus_areas": [
            "operational excellence and execution",
            "cross-functional collaboration",
            "organizational impact",
            "process improvement",
            "scaling teams and systems",
            "stakeholder management",
            "strategic planning and execution",
            "performance metrics and KPIs",
            "change management",
            "building and maintaining partnerships"
        ],
        "question_types": {"strategic": 40, "leadership": 35, "operational": 25},
        "tone": "operational, focused on execution and organizational impact",
        "example_questions": [
            "How have you contributed to improving processes in your organization?",
            "Tell me about leading a cross-functional initiative.",
            "How do you ensure alignment between different teams?",
            "Describe a time you had to drive organizational change.",
            "How do you measure and improve team performance?"
        ]
    },
    "Vice President": {
        "title": "Vice President",
        "focus_areas": [
            "strategic leadership and direction",
            "building and scaling teams",
            "cross-departmental influence",
            "OKRs and goal setting",
            "executive communication",
            "talent development and retention",
            "portfolio management",
            "stakeholder relationships",
            "innovation and transformation",
            "representing the organization"
        ],
        "question_types": {"leadership": 45, "strategic": 35, "behavioral": 20},
        "tone": "executive, focused on leadership, influence, and strategic impact",
        "example_questions": [
            "How do you build and maintain high-performing teams?",
            "Tell me about a strategic initiative you led from concept to execution.",
            "How do you handle competing priorities across departments?",
            "Describe your approach to developing talent.",
            "How do you communicate complex technical concepts to non-technical executives?"
        ]
    }
}

# Difficulty levels with detailed expectations
DIFFICULTY_PROFILES = {
    "intern": {
        "experience_range": "0-1 years / Student",
        "question_count": 12,
        "technical_depth": "fundamental",
        "expectations": [
            "Basic understanding of core concepts",
            "Eagerness to learn and grow",
            "Academic projects and coursework",
            "Problem-solving approach (not necessarily optimal solutions)",
            "Communication of thought process",
            "Enthusiasm and curiosity"
        ],
        "avoid": [
            "Complex system design questions",
            "Leadership and mentoring questions",
            "Questions requiring production experience",
            "Deep architectural trade-offs"
        ],
        "focus": "Learning ability, potential, basic fundamentals, enthusiasm"
    },
    "junior": {
        "experience_range": "1-2 years",
        "question_count": 15,
        "technical_depth": "foundational with some practical application",
        "expectations": [
            "Solid grasp of fundamentals",
            "Some real-world project experience",
            "Basic debugging and troubleshooting",
            "Understanding of development workflows",
            "Collaboration in team settings",
            "Growth mindset and coachability"
        ],
        "avoid": [
            "Senior-level architecture questions",
            "Questions about leading large teams",
            "Complex distributed systems"
        ],
        "focus": "Core skills, problem-solving, learning from experiences, growth potential"
    },
    "associate": {
        "experience_range": "2-4 years",
        "question_count": 18,
        "technical_depth": "solid practical knowledge",
        "expectations": [
            "Independent task completion",
            "Moderate complexity problem solving",
            "Code quality and best practices awareness",
            "Effective collaboration",
            "Some mentoring of juniors",
            "Project ownership for defined scope"
        ],
        "avoid": [
            "Questions requiring 10+ years experience",
            "C-level strategic questions"
        ],
        "focus": "Technical competence, ownership, collaboration, growing independence"
    },
    "senior": {
        "experience_range": "5-8 years",
        "question_count": 20,
        "technical_depth": "deep expertise with architectural thinking",
        "expectations": [
            "Technical leadership and decision making",
            "System design and architecture",
            "Mentoring and code reviews",
            "Cross-team collaboration",
            "Trade-off analysis",
            "Business impact awareness",
            "Production systems experience",
            "Handling ambiguity"
        ],
        "avoid": [
            "Basic/trivial technical questions",
            "Questions appropriate for entry-level"
        ],
        "focus": "Architecture, trade-offs, leadership, business impact, mentoring"
    },
    "lead": {
        "experience_range": "8+ years",
        "question_count": 22,
        "technical_depth": "expert with strategic vision",
        "expectations": [
            "Technical strategy and vision",
            "Team building and development",
            "Cross-organizational influence",
            "Executive communication",
            "Complex system ownership",
            "Innovation and improvement",
            "Risk management",
            "Stakeholder management"
        ],
        "avoid": [
            "Entry-level technical questions",
            "Questions not befitting seniority"
        ],
        "focus": "Strategy, leadership, organizational impact, technical vision"
    }
}

# Role-specific technical focus areas (inferred from resume/skills)
ROLE_TECHNICAL_FOCUS = {
    "software_engineer": {
        "keywords": ["software", "developer", "engineer", "programming", "coding", "full stack", "backend", "frontend"],
        "technical_areas": [
            "coding and algorithms",
            "system design",
            "API design",
            "database design",
            "testing strategies",
            "CI/CD and DevOps",
            "code review",
            "debugging production issues"
        ]
    },
    "data_scientist": {
        "keywords": ["data scientist", "machine learning", "ML", "AI", "analytics", "statistical"],
        "technical_areas": [
            "machine learning algorithms",
            "statistical analysis",
            "data preprocessing and cleaning",
            "model evaluation and validation",
            "feature engineering",
            "A/B testing",
            "data visualization",
            "production ML systems"
        ]
    },
    "data_engineer": {
        "keywords": ["data engineer", "ETL", "pipeline", "data warehouse", "big data"],
        "technical_areas": [
            "data pipeline design",
            "ETL processes",
            "data warehouse architecture",
            "big data technologies",
            "data quality and governance",
            "real-time vs batch processing",
            "data modeling",
            "performance optimization"
        ]
    },
    "devops_engineer": {
        "keywords": ["devops", "SRE", "infrastructure", "cloud", "kubernetes", "docker"],
        "technical_areas": [
            "CI/CD pipeline design",
            "infrastructure as code",
            "container orchestration",
            "monitoring and alerting",
            "incident response",
            "security best practices",
            "cost optimization",
            "high availability design"
        ]
    },
    "product_manager": {
        "keywords": ["product manager", "product owner", "PM", "product"],
        "technical_areas": [
            "product strategy and roadmap",
            "user research and feedback",
            "prioritization frameworks",
            "metrics and KPIs",
            "stakeholder management",
            "agile methodologies",
            "go-to-market strategy",
            "competitive analysis"
        ]
    },
    "default": {
        "keywords": [],
        "technical_areas": [
            "problem-solving approach",
            "technical decision making",
            "collaboration and communication",
            "project execution",
            "quality and attention to detail"
        ]
    }
}


def detect_role_category(role: str, skills: List[str]) -> str:
    """
    Detect the role category based on role title and skills from resume.
    This helps generate more relevant technical questions.
    """
    role_lower = (role or "").lower()
    skills_lower = " ".join(skills).lower() if skills else ""
    combined = f"{role_lower} {skills_lower}"

    for category, config in ROLE_TECHNICAL_FOCUS.items():
        if category == "default":
            continue
        for keyword in config["keywords"]:
            if keyword.lower() in combined:
                return category

    return "default"


def get_interviewer_profile(interviewer: str) -> Dict[str, Any]:
    """Get the profile for an interviewer, with fallback to Manager profile."""
    return INTERVIEWER_PROFILES.get(interviewer, INTERVIEWER_PROFILES["Manager"])


def get_difficulty_profile(difficulty: str) -> Dict[str, Any]:
    """Get the profile for a difficulty level, with fallback to junior."""
    lvl = (difficulty or "").lower()
    if lvl in DIFFICULTY_PROFILES:
        return DIFFICULTY_PROFILES[lvl]
    # Map variations
    if lvl in ["mid", "middle", "intermediate"]:
        return DIFFICULTY_PROFILES["associate"]
    if lvl in ["staff", "principal", "architect"]:
        return DIFFICULTY_PROFILES["lead"]
    return DIFFICULTY_PROFILES["junior"]


def count_for_level(level: str) -> int:
    """Return question count based on difficulty level."""
    profile = get_difficulty_profile(level)
    return profile.get("question_count", 15)


def calculate_question_distribution(interviewers: List[str], total_questions: int) -> Dict[str, int]:
    """
    Calculate how many questions each interviewer should ask.
    Distributes questions realistically based on interviewer roles.
    """
    if not interviewers:
        return {"Interviewer": total_questions}

    # Weight different interviewers (Manager and Tech Lead typically ask more)
    weights = {
        "Manager": 3,
        "Tech Lead": 3,
        "HR": 2,
        "CEO": 1,
        "CFO": 1,
        "President": 1,
        "Vice President": 1
    }

    total_weight = sum(weights.get(i, 2) for i in interviewers)
    distribution = {}

    remaining = total_questions - 3  # Reserve 3 for warmup (HR typically does these)

    for interviewer in interviewers:
        weight = weights.get(interviewer, 2)
        count = max(1, int((weight / total_weight) * remaining))
        distribution[interviewer] = count

    # Adjust to match total (add remaining to Manager or first interviewer)
    allocated = sum(distribution.values())
    diff = remaining - allocated
    if diff != 0:
        primary = "Manager" if "Manager" in distribution else interviewers[0]
        distribution[primary] = distribution.get(primary, 0) + diff

    return distribution


def questions_schema_text() -> str:
    """Return the JSON schema for question generation."""
    return """
Return ONLY JSON with this shape:
{
  "questions": [
    {
      "prompt": string (the actual question to ask),
      "topic": string (e.g., "System Design", "Behavioral", "Technical", "Leadership"),
      "interviewer": string (who asks this question),
      "type": string (one of: "warmup", "technical", "behavioral", "situational", "strategic", "analytical"),
      "idealAnswer": string (a strong example answer for reference),
      "rubric": {
        "content": string (what makes a good answer content-wise),
        "clarity": string (expectations for communication clarity),
        "structure": string (expected answer structure, e.g., STAR method)
      }
    }
  ]
}
No extra keys. No prose outside JSON. No markdown.
""".strip()


def build_questions_system_prompt() -> str:
    """Build the system prompt for question generation."""
    return """You are an expert interview question generator with 20+ years of experience conducting and designing technical and behavioral interviews at top companies like Google, Amazon, Microsoft, and Meta.

Your task is to generate realistic, high-quality interview questions that:
1. Are specific to the candidate's background and the role they're applying for
2. Match the difficulty level appropriately
3. Reflect what each interviewer type would realistically ask
4. Include a mix of technical, behavioral, and situational questions
5. Have practical, real-world relevance (not textbook questions)

CRITICAL RULES:
- Each question must be unique and not repeat concepts with different wording
- Questions must be calibrated to the difficulty level
- Technical questions should reference specific technologies from the candidate's skills
- Behavioral questions should use "Tell me about a time..." or similar formats
- Include the interviewer's name/role in each question object
- Provide detailed ideal answers and rubrics

You must return STRICT JSON ONLY following the given schema. No prose, no markdown, no explanations."""


def build_questions_user_prompt(
    role: str,
    difficulty: str,
    interviewers: List[str],
    requested_total: int,
    skills: List[str],
    keywords: List[str],
    job_description: Optional[str] = None,
) -> str:
    """
    Build a comprehensive user prompt for generating interview questions.
    This creates role-specific, interviewer-appropriate, difficulty-calibrated questions.
    """
    # Get profiles
    difficulty_profile = get_difficulty_profile(difficulty)
    role_category = detect_role_category(role, skills)
    role_tech_focus = ROLE_TECHNICAL_FOCUS.get(role_category, ROLE_TECHNICAL_FOCUS["default"])

    # Calculate question count
    total_allowed = difficulty_profile["question_count"]
    target = min(requested_total or total_allowed, total_allowed)

    # Set default interviewers if none provided
    if not interviewers:
        interviewers = ["HR", "Manager", "Tech Lead"]

    # Calculate question distribution per interviewer
    distribution = calculate_question_distribution(interviewers, target)

    # Build interviewer section with detailed guidance
    interviewer_sections = []
    for interviewer in interviewers:
        profile = get_interviewer_profile(interviewer)
        count = distribution.get(interviewer, 2)

        section = f"""
### {interviewer} ({profile['title']}) - {count} questions
Focus Areas: {', '.join(profile['focus_areas'][:5])}
Question Type Mix: {', '.join(f"{k}: {v}%" for k, v in profile['question_types'].items())}
Tone: {profile['tone']}
Example Questions:
{chr(10).join(f"  - {q}" for q in profile['example_questions'][:3])}
"""
        interviewer_sections.append(section)

    # Build skills section
    skills_text = ", ".join(skills[:30]) if skills else "Not specified"
    keywords_text = ", ".join(keywords[:20]) if keywords else "Not specified"
    tech_focus = ", ".join(role_tech_focus["technical_areas"][:6])

    # Build job description section if provided
    jd_section = ""
    jd_text = (job_description or "").strip()
    if len(jd_text) >= 40:
        jd_section = f"""
## JOB DESCRIPTION ALIGNMENT (HIGH PRIORITY)
{jd_text[:5000]}

IMPORTANT JD REQUIREMENTS:
- Generate questions that directly verify the candidate can perform JD responsibilities
- Include questions about specific tools/technologies mentioned in the JD
- Ask about experiences relevant to the job requirements
- Include 2-3 questions probing potential gaps (skills in JD but not in resume)
"""

    # Build the comprehensive prompt
    prompt = f"""
# INTERVIEW QUESTION GENERATION REQUEST

## CANDIDATE PROFILE
- **Target Role**: {role or "Software Professional"}
- **Experience Level**: {difficulty or "Junior"} ({difficulty_profile['experience_range']})
- **Core Skills**: {skills_text}
- **Keywords/Domains**: {keywords_text}
- **Detected Role Category**: {role_category.replace('_', ' ').title()}
- **Technical Focus Areas**: {tech_focus}

{jd_section}

## DIFFICULTY CALIBRATION: {difficulty.upper() if difficulty else "JUNIOR"}
**Experience Range**: {difficulty_profile['experience_range']}
**Technical Depth Expected**: {difficulty_profile['technical_depth']}

**What to Expect from This Level**:
{chr(10).join(f"- {exp}" for exp in difficulty_profile['expectations'])}

**Avoid These for This Level**:
{chr(10).join(f"- {avoid}" for avoid in difficulty_profile['avoid'])}

**Focus On**: {difficulty_profile['focus']}

## INTERVIEWER PANEL
{chr(10).join(interviewer_sections)}

## QUESTION GENERATION REQUIREMENTS

### Total Questions: {target}

### Question Structure:
1. **Warmup Questions (First 3)** - Asked by HR or Manager
   - Start with a friendly, open-ended introduction question
   - Ask about motivation for applying
   - Ask about self-assessment (strengths/growth areas)

2. **Technical Questions** - Based on skills and role
   - Ask about specific technologies from the candidate's skill set: {skills_text}
   - Include practical scenarios, not just theoretical questions
   - Calibrate complexity to {difficulty} level

3. **Behavioral Questions** - Using STAR method prompts
   - "Tell me about a time when..."
   - "Describe a situation where..."
   - "Give me an example of..."

4. **Situational Questions** - Hypothetical scenarios
   - "How would you handle..."
   - "What would you do if..."
   - "Imagine you're faced with..."

### Quality Requirements:
- Each question must be UNIQUE (no conceptual duplicates)
- Questions must reference SPECIFIC skills from the candidate's profile
- Technical questions should mention actual technologies: {skills_text}
- Behavioral questions should be role-relevant
- Ideal answers should be detailed (3-5 sentences minimum)
- Rubrics should give clear evaluation criteria

### Distribution by Interviewer:
{chr(10).join(f"- {name}: {count} questions" for name, count in distribution.items())}

## OUTPUT FORMAT
{questions_schema_text()}

Generate exactly {target} high-quality, realistic interview questions now.
"""

    return prompt.strip()



# ---------- Response schema (validates the LLM output) ----------
class Rare(BaseModel):
    readability: float = Field(ge=0, le=5)
    applicability: float = Field(ge=0, le=5)
    remarkability: float = Field(ge=0, le=5)
    total: float = Field(ge=0, le=5)

class ParseOut(BaseModel):
    skills: List[str]
    fallbackRoles: List[str]
    rare: Rare
    atsScore: int = Field(ge=0, le=100)
    atsSuggestions: List[str]
    keywords: List[str] = []

    # NEW: ATS breakdown + JD match fields (optional so resume-only still works)
    atsBreakdown: Optional[Dict[str, Any]] = None

    # JD-mode extras (optional)
    jobMatchScore: Optional[int] = None
    matchedKeywords: Optional[List[str]] = []
    missingKeywords: Optional[List[str]] = []
    hasJobDescription: Optional[bool] = False




# ---------- OpenAI call with strict JSON ----------
def call_openai(resume_text: str, job_desc_text: Optional[str] = None) -> Dict[str, Any]:
    resume_snippet = (resume_text or "")[:28000]
    jd_snippet = (job_desc_text or "").strip()[:12000]

    has_jd = bool(job_desc_text and len(job_desc_text.strip()) >= 40)

    system_prompt = SYSTEM_PROMPT_JD if has_jd else SYSTEM_PROMPT
    user_content = (
        f"RESUME TEXT:\n\n{resume_snippet}\n\nJOB DESCRIPTION TEXT:\n\n{jd_snippet}"
        if has_jd
        else f"RESUME TEXT:\n\n{resume_snippet}"
    )

    # 1) Call OpenAI and parse JSON
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.3,
        presence_penalty=0.2,
        frequency_penalty=0.2,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
    )

    raw_json = resp.choices[0].message.content
    data = json.loads(raw_json)

        # --- Normalize/repair model output for Rare if it comes back malformed ---
    rare_val = data.get("rare")

    if isinstance(rare_val, (int, float)):
        # Model incorrectly returned rare as a number, convert to structured object
        v = float(rare_val)
        data["rare"] = {
            "readability": v,
            "applicability": v,
            "remarkability": v,
            "total": v,
        }
    elif isinstance(rare_val, dict):
        # Ensure required keys exist
        r = rare_val
        data["rare"] = {
            "readability": float(r.get("readability", 4.0) or 4.0),
            "applicability": float(r.get("applicability", 4.0) or 4.0),
            "remarkability": float(r.get("remarkability", 4.0) or 4.0),
            "total": float(r.get("total", 0.0) or 0.0),
        }
    else:
        # Missing or wrong type
        data["rare"] = {
            "readability": 4.0,
            "applicability": 4.0,
            "remarkability": 4.0,
            "total": 4.0,
        }


    skills_val = data.get("skills")
    if isinstance(skills_val, str):
        data["skills"] = [s.strip() for s in skills_val.split(",") if s.strip()]
    elif not isinstance(skills_val, list):
        data["skills"] = []



    # 2) Validate to ParseOut -> NOW parsed exists
    parsed = ParseOut.model_validate(data).model_dump()
    parsed["hasJobDescription"] = has_jd

    # 3) Apply deterministic ATS override (Jobscan/Enhancv style)
    ats = compute_ats_score(resume_text, job_desc_text if has_jd else None)

    parsed["atsScore"] = int(ats.get("atsScore", parsed.get("atsScore", 0)) or 0)
    parsed["atsBreakdown"] = ats.get("breakdown")

    parsed["jobMatchScore"] = ats.get("jobMatchScore")
    parsed["matchedKeywords"] = ats.get("matchedKeywords", [])
    parsed["missingKeywords"] = ats.get("missingKeywords", [])

    # 4) Safety defaults
    if not parsed.get("fallbackRoles"):
        parsed["fallbackRoles"] = ["Data Scientist", "ML Engineer", "Python Developer", "Business Analyst"]

    r = parsed.get("rare") or {}
    if not isinstance(r.get("total"), (int, float)):
        avg = (float(r["readability"]) + float(r["applicability"]) + float(r["remarkability"])) / 3
        parsed["rare"]["total"] = round(avg, 1)

    if has_jd and parsed.get("jobMatchScore") is None:
        parsed["jobMatchScore"] = parsed["atsScore"]

    return parsed


# ---------- Parse Resume ----------
@app.post("/api/parse-resume")
async def parse_resume(
    file: UploadFile = File(...),
    job_description: Optional[str] = Form(None),
) -> Dict[str, Any]:
    data = await file.read()
    name = (file.filename or "").lower()

    # --- Extract text from file ---
    if name.endswith(".pdf"):
        text = extract_pdf_text(data)
    elif name.endswith(".docx"):
        text = extract_docx_text(data)
    else:
        try:
            text = extract_pdf_text(data)
        except Exception:
            try:
                text = extract_docx_text(data)
            except Exception:
                return {"error": "Unsupported file. Please upload a PDF or DOCX."}

    text = clean_text(text)

    if len(text) < 20:
        return {"error": "Could not read text from file (image-only PDF?). Try a text-based PDF/DOCX."}

    # --- Optional JD mode ---
    jd_text = (job_description or "").strip()
    has_jd = len(jd_text) >= 40  # threshold to avoid accidental short inputs

    try:
        # IMPORTANT: call_openai must support the optional second argument
        # Mode A: call_openai(resume_text)
        # Mode B: call_openai(resume_text, job_desc_text)
        result = call_openai(text, jd_text if has_jd else None)

        # attach a flag so frontend can show "JD match mode enabled"
        if isinstance(result, dict):
            result["hasJobDescription"] = has_jd

        # --- Save to Postgres (Supabase) ---
        try:
            raw_json = json.dumps(result, ensure_ascii=False)
        except Exception:
            raw_json = "{}"

        ats_score = int(result.get("atsScore", 0) or 0)
        rare = result.get("rare") or {}
        rare_total = float(rare.get("total", 0.0) or 0.0)
        skills = result.get("skills") or []
        keywords = result.get("keywords") or []

        with Session(engine) as session:
            db_row = Resume(
                original_filename=file.filename or "unknown",
                ats_score=ats_score,
                rare_total=rare_total,
                skills=skills,
                keywords=keywords,
                raw_json=raw_json,
            )
            session.add(db_row)
            session.commit()

        return result

    except Exception as e:
        print("OpenAI error:", e)

        # fallback JSON (same as before)
        fallback = {
            "skills": [],
            "fallbackRoles": ["Data Scientist", "ML Engineer", "Python Developer", "Business Analyst"],
            "rare": {"readability": 4.5, "applicability": 4.5, "remarkability": 4.5, "total": 4.5},
            "atsScore": 85,
            "atsSuggestions": [
                "Add role-specific keywords matching the target JD.",
                "Quantify results with % and $ metrics.",
                "Use a simple one-column layout without tables."
            ],
            "keywords": [],
            "hasJobDescription": has_jd,
        }

        try:
            raw_json = json.dumps(fallback, ensure_ascii=False)
            with Session(engine) as session:
                db_row = Resume(
                    original_filename=file.filename or "unknown",
                    ats_score=fallback["atsScore"],
                    rare_total=fallback["rare"]["total"],
                    skills=fallback["skills"],
                    keywords=fallback["keywords"],
                    raw_json=raw_json,
                )
                session.add(db_row)
                session.commit()
        except Exception as db_err:
            print("Failed to store fallback resume:", db_err)

        return fallback


# ---------- Transcribe Audio (Speech-to-Text via OpenAI Whisper) ----------
@app.post("/api/transcribe")
async def transcribe_audio(file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    Accepts an audio file (webm/opus/etc from MediaRecorder),
    sends it to OpenAI Whisper, and returns { transcript: "..." }.
    """
    try:
        audio_bytes = await file.read()
        if not audio_bytes:
            raise HTTPException(status_code=400, detail="Empty audio file")

        # Write to temp file because OpenAI Python SDK expects a file-like object
        with tempfile.NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name

        try:
            with open(tmp_path, "rb") as f:
                stt_resp = client.audio.transcriptions.create(
                    model="whisper-1",   # we can switch to another model, later after MVP.
                    file=f,
                    response_format="json",
                    temperature=0,
                    language="en"
                )

            text = (stt_resp.text or "").strip()
            return {"transcript": text}
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    except HTTPException:
        raise
    except Exception as e:
        logging.exception("OpenAI transcription failed")
        raise HTTPException(status_code=500, detail=f"Transcription failed: {e}")


# ---------- Generate Questions ----------
class Question(BaseModel):
    id: int
    prompt: str
    topic: str = "Mixed"
    interviewer: str = "Interviewer"
    type: str = "technical"
    idealAnswer: Optional[str] = None
    rubric: Optional[Dict[str, Any]] = None

class GenerateReq(BaseModel):
    role: str
    difficulty: str
    interviewers: List[str] = []
    count: int = 20
    skills: List[str] = []
    keywords: List[str] = []
    jobDescription: Optional[str] = None

class GenerateResp(BaseModel):
    meta: Dict[str, Any]
    questions: List[Question]

@app.post("/api/generate-questions", response_model=GenerateResp)
def generate_questions(req: GenerateReq):
    logging.info("REQ /api/generate-questions: %s", req.dict())

    target_total = min(req.count or count_for_level(req.difficulty), count_for_level(req.difficulty))

    try:
        system = build_questions_system_prompt()
        user = build_questions_user_prompt(
            req.role, req.difficulty, req.interviewers, target_total, req.skills or [], req.keywords or [], req.jobDescription,
            )

        r = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.8,
            presence_penalty=0.2,
            frequency_penalty=0.2,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )

        content = r.choices[0].message.content
        logging.info("RAW LLM JSON (trunc): %s", content[:900])

        data = json.loads(content)
        raw_list = data.get("questions", [])
        if not isinstance(raw_list, list) or not raw_list:
            raise ValueError("Model did not return 'questions' list")

        seen = set()
        cleaned: List[Question] = []

        for idx, item in enumerate(raw_list, start=1):
            if isinstance(item, str):
                item_obj = {"prompt": item}
            elif isinstance(item, dict):
                item_obj = item
            else:
                continue

            prompt = (item_obj.get("prompt") or "").strip()
            if not prompt or prompt in seen:
                continue
            seen.add(prompt)

            interviewer = item_obj.get("interviewer")
            if not isinstance(interviewer, str) or not interviewer.strip():
                if req.interviewers:
                    interviewer = req.interviewers[(idx - 1) % max(1, len(req.interviewers))]
                else:
                    interviewer = "Interviewer"

            q_type = item_obj.get("type")
            if not isinstance(q_type, str) or not q_type.strip():
                q_type = "warmup" if idx <= 3 else "technical"

            topic = item_obj.get("topic")
            if not isinstance(topic, str) or not topic.strip():
                topic = "Mixed"

            ideal = item_obj.get("idealAnswer")
            if isinstance(ideal, (list, dict)):
                ideal = json.dumps(ideal, ensure_ascii=False)
            elif not isinstance(ideal, str):
                ideal = None

            rubric = item_obj.get("rubric")
            if not isinstance(rubric, dict):
                rubric = {
                    "content": "Accuracy and relevance to the prompt; includes key concepts.",
                    "clarity": "Clear, concise, minimal filler; easy to follow.",
                    "structure": "Logical flow (context → approach → result); concrete examples."
                }

            q = Question(
                id=idx,
                prompt=prompt,
                topic=topic,
                interviewer=interviewer,
                type=q_type,
                idealAnswer=ideal,
                rubric=rubric,
            )
            cleaned.append(q)

            if len(cleaned) >= target_total:
                break

        resp = GenerateResp(
            meta={"role": req.role, "difficulty": req.difficulty, "questionCount": len(cleaned)},
            questions=cleaned
        )
        logging.info("RESP /api/generate-questions: %s", resp.model_dump())
        return resp

    except Exception as e:
        logging.exception("generate-questions failed")
        raise HTTPException(status_code=500, detail=f"Failed to generate questions: {e}")


# ---------- Scoring: end-of-interview feedback ----------
class AnswerIn(BaseModel):
    id: int
    prompt: str
    interviewer: str
    type: str
    userAnswer: str
    idealAnswer: Optional[str] = None

class QuestionFeedback(BaseModel):
    id: int
    prompt: str
    interviewer: str
    type: str
    userAnswer: str
    idealAnswer: str
    scores: Dict[str, float]
    strengths: List[str]
    improvements: List[str]
    suggestedAnswer: str

class OverallFeedback(BaseModel):
    overallScore: float
    summary: str
    strengths: List[str]
    improvements: List[str]

class ScoreInterviewReq(BaseModel):
    role: str
    difficulty: str
    answers: List[AnswerIn]
    plan: dict | None = None

class ScoreInterviewResp(BaseModel):
    meta: Dict[str, Any]
    questions: List[QuestionFeedback]
    overall: OverallFeedback


# -------------------- Interview Answer Models --------------------
# These models handle the submission of individual interview answers
# for the two-way communication feature (Phase 1)

class SubmitAnswerRequest(BaseModel):
    """
    Request model for submitting an interview answer.
    Validates all required fields for storing a Q&A pair.
    """
    session_id: str = Field(..., description="UUID of the interview session")
    question_id: int = Field(..., description="The order/number of the question (1-indexed)")
    question_text: str = Field(..., description="The actual question that was asked")
    question_intent: str = Field(..., description="What the question was testing (e.g., 'technical skills', 'problem solving')")
    role: str = Field(..., description="The role being interviewed for (e.g., 'Software Engineer')")
    user_answer: str = Field(..., description="The user's transcribed response")
    transcript_raw: Optional[str] = Field(None, description="Full transcript with timestamps if available")
    audio_duration_seconds: Optional[float] = Field(None, description="Duration of the audio response in seconds")


class SubmitAnswerResponse(BaseModel):
    """
    Response model for a successfully submitted answer.
    """
    success: bool
    answer_id: str
    message: str


class AnswerDetail(BaseModel):
    """
    Response model for a single answer in the list.
    Contains all stored answer data for display/processing.
    """
    id: str
    session_id: str
    question_id: int
    question_text: str
    question_intent: str
    role: str
    user_answer: str
    transcript_raw: Optional[str]
    audio_duration_seconds: Optional[float]
    answer_timestamp: datetime
    created_at: datetime


class RetrieveAnswersResponse(BaseModel):
    """
    Response model for retrieving all answers in a session.
    """
    success: bool
    session_id: str
    total_answers: int
    answers: List[AnswerDetail]


# -------------------- Interview Session Models --------------------
# These models handle creating a new interview session at the START
# so answers can be saved in real-time during the interview

class CreateSessionRequest(BaseModel):
    """
    Request model for creating a new interview session.
    Contains the interview configuration/metadata.
    """
    role: str = Field(..., description="The role being interviewed for (e.g., 'Software Engineer')")
    difficulty: str = Field(..., description="Interview difficulty level (e.g., 'Junior', 'Senior')")
    question_count: int = Field(..., description="Total number of questions in this interview")
    interviewer_names: List[str] = Field(default=[], description="List of interviewer names/roles")
    plan: Optional[Dict[str, Any]] = Field(None, description="Full interview plan with questions")


class CreateSessionResponse(BaseModel):
    """
    Response model for a successfully created session.
    Returns the session_id needed for submitting answers.
    """
    success: bool
    session_id: str
    message: str


class SessionDetailResponse(BaseModel):
    """
    Response model for fetching full session details.
    Includes session metadata, plan, and all answers.
    Used by Feedback page to retrieve interview data from database.
    """
    success: bool
    session_id: str
    role: str
    difficulty: str
    question_count: int
    interviewer_names: List[str]
    plan: Optional[Dict[str, Any]]
    answers: List[AnswerDetail]
    created_at: datetime


# -------------------- Semantic Search Models --------------------
# These models handle semantic search for finding similar answers
# based on embedding similarity within an interview session

class SearchAnswersRequest(BaseModel):
    """
    Request model for semantic search across interview answers.
    Searches within a specific session using embedding similarity.
    """
    session_id: str = Field(..., description="UUID of the interview session to search within")
    query: str = Field(..., description="The search query text (e.g., 'tell me about machine learning')")
    top_k: int = Field(default=5, ge=1, le=10, description="Number of top similar answers to return (1-10)")


class SearchAnswerResult(BaseModel):
    """
    A single search result containing the answer and its similarity score.
    """
    answer_id: str = Field(..., description="UUID of the matched answer")
    question_id: int = Field(..., description="The question number in the interview")
    question_text: str = Field(..., description="The question that was asked")
    user_answer: str = Field(..., description="The user's response")
    role: str = Field(..., description="The interviewer role who asked this question")
    similarity_score: float = Field(..., description="Cosine similarity score (0-1, higher = more similar)")


class SearchAnswersResponse(BaseModel):
    """
    Response model for semantic search results.
    Returns answers sorted by similarity (highest first).
    """
    success: bool
    session_id: str
    query: str
    total_results: int
    results: List[SearchAnswerResult]


# -------------------- Phase 1.3: Intelligent Interview Flow Models --------------------
# These models handle the intelligent question generation endpoints

class NextQuestionRequest(BaseModel):
    """Request model for getting the next contextually-aware question."""
    session_id: str = Field(..., description="UUID of the interview session")
    current_question_number: int = Field(..., gt=0, description="Current question number (1-indexed)")
    role: str = Field(..., description="Job role being interviewed for")
    difficulty: str = Field(..., description="Interview difficulty level")
    total_questions: int = Field(default=10, gt=0, description="Total planned questions")


class QuestionData(BaseModel):
    """The generated question data."""
    id: int = Field(..., description="Question number/ID")
    text: str = Field(..., description="The question text")
    intent: str = Field(..., description="Question intent category")
    type: str = Field(..., description="Question type (standard/follow_up/challenge/deep_dive/reference)")


class QuestionReferences(BaseModel):
    """References to past answers if the question builds on previous context."""
    question_id: Optional[int] = Field(None, description="ID of referenced question")
    excerpt: Optional[str] = Field(None, description="Excerpt from referenced answer")


class QuestionMetadata(BaseModel):
    """Metadata about the question generation decision."""
    decision_reason: str = Field(..., description="Why this question type was chosen")
    patterns_detected: List[str] = Field(default=[], description="Patterns found in conversation")
    conversation_stage: str = Field(..., description="Stage of interview (early/mid/late)")
    action_taken: str = Field(..., description="Action type taken")
    context_used: Optional[str] = Field(None, description="Context used for generation")
    generated_at: Optional[str] = Field(None, description="ISO timestamp of generation")


class NextQuestionResponse(BaseModel):
    """Response model for the next question endpoint."""
    question: QuestionData
    interviewer_comment: Optional[str] = Field(None, description="Pre-question comment")
    references: QuestionReferences
    metadata: QuestionMetadata


class SubmitAndNextRequest(BaseModel):
    """Request model for submitting an answer and getting the next question."""
    session_id: str = Field(..., description="UUID of the interview session")
    question_id: int = Field(..., gt=0, description="The question number being answered")
    question_text: str = Field(..., description="The question that was asked")
    question_intent: str = Field(..., description="Intent category of the question")
    role: str = Field(..., description="Job role being interviewed for")
    user_answer: str = Field(..., description="The user's transcribed response")
    transcript_raw: Optional[str] = Field(None, description="Full transcript with timestamps")
    audio_duration_seconds: Optional[float] = Field(None, ge=0, description="Audio duration in seconds")
    difficulty: str = Field(..., description="Interview difficulty level")
    total_questions: int = Field(default=10, gt=0, description="Total planned questions")


class SubmitAndNextResponse(BaseModel):
    """Response model for submit-and-next endpoint."""
    answer_stored: bool = Field(..., description="Whether the answer was stored")
    answer_id: Optional[str] = Field(None, description="UUID of the stored answer")
    embedding_generated: Optional[bool] = Field(None, description="Whether embedding was generated")
    next_question: NextQuestionResponse


class ContradictionDetail(BaseModel):
    """Details about a detected contradiction."""
    previous_statement: Optional[str] = Field(None, description="Earlier contradicting statement")
    current_statement: Optional[str] = Field(None, description="Current contradicting statement")
    contradiction_type: Optional[str] = Field(None, description="Type of contradiction")
    confidence_score: Optional[float] = Field(None, description="Confidence score (0-1)")
    explanation: Optional[str] = Field(None, description="Explanation of the contradiction")


class QualityMetrics(BaseModel):
    """Quality metrics for the conversation."""
    avg_answer_length: float = Field(..., description="Average word count of answers")
    star_format_usage: float = Field(..., description="Ratio of answers using STAR format")
    specificity_score: float = Field(..., description="Average specificity score")


class ConversationStateResponse(BaseModel):
    """Response model for conversation state analysis."""
    total_answers: int = Field(..., description="Number of answers in the session")
    topics_discussed: List[str] = Field(default=[], description="Topics mentioned in answers")
    repeated_topics: Dict[str, int] = Field(default={}, description="Topics with mention counts")
    contradictions_detected: List[ContradictionDetail] = Field(default=[], description="Detected contradictions")
    conversation_summary: str = Field(..., description="Summary of the conversation")
    quality_metrics: QualityMetrics
    recommendations: List[str] = Field(default=[], description="Recommendations for interviewer")


class PatternsResponse(BaseModel):
    """Response model for detected patterns in conversation."""
    contradictions: List[ContradictionDetail] = Field(default=[], description="Detected contradictions")
    repeated_topics: Dict[str, int] = Field(default={}, description="Topics with mention counts")
    weak_answers: List[int] = Field(default=[], description="Question IDs with weak answers")
    strong_areas: List[str] = Field(default=[], description="Topics where candidate is strong")
    gaps: List[str] = Field(default=[], description="Expected topics not yet covered")


def build_scoring_system_prompt() -> str:
    return (
        "You are a supportive but honest interview coach evaluating a candidate's answers.\n\n"
        "You will receive JSON with:\n"
        "{\n"
        '  \"role\": string,\n'
        '  \"difficulty\": string,\n'
        '  \"answers\": [\n'
        "    {\n"
        '      \"id\": number,\n'
        '      \"prompt\": string,\n'
        '      \"interviewer\": string,\n'
        '      \"type\": string,\n'
        '      \"userAnswer\": string,\n'
        '      \"idealAnswer\": string | null\n'
        "    }\n"
        "  ]\n"
        "}\n\n"
        "Calibration rules:\n"
        "- Use BOTH role and difficulty to set the bar.\n"
        "- For SENIOR: expect ownership, trade-offs, metrics, and impact. Vague answers must NOT get high scores.\n"
        "- For JUNIOR/INTERN: focus more on fundamentals, clarity and learning mindset.\n"
        "- 5.0 should be rare and outstanding. 3.0 is average. 2.0 or below is weak.\n\n"
        "For EACH answer, produce:\n"
        "- scores (0–5 floats) for: content, structure, clarity, confidence, relevance.\n"
        "  * content: correctness, depth, use of appropriate tools/techniques for the role.\n"
        "  * structure: logical flow (context → actions → result).\n"
        "  * clarity: understandable, avoids rambling, uses precise language.\n"
        "  * confidence: steady, assertive wording vs. very unsure or self-contradictory.\n"
        "  * relevance: how directly it answers the prompt.\n"
        "- strengths: 2–4 short bullet-style strings that highlight what they did WELL.\n"
        "- improvements: 2–4 short bullet-style strings that are SPECIFIC and ACTIONABLE.\n"
        "  * Reference missing metrics, tools, steps, or examples.\n"
        "  * Avoid generic advice like \"be more confident\" unless you also say HOW.\n"
        "- suggestedAnswer: a concise, strong answer tailored to this candidate and this role.\n\n"
        "Overall section:\n"
        "- overallScore: single float 0–5 summarizing the interview.\n"
        "- summary: 2–3 sentences, honest but encouraging.\n"
        "- strengths: 3–5 positive bullet-style observations.\n"
        "- improvements: 3–5 very practical next steps (what to practice, how to answer better).\n\n"
        "Tone guidelines:\n"
        "- Always kind, constructive, and growth-oriented.\n"
        "- Do NOT inflate scores if the answer is vague or off-topic.\n"
        "- Focus on what they can DO next time to sound like a stronger {role}.\n\n"
        "You MUST return STRICT JSON ONLY with EXACTLY this shape:\n"
        "{\n"
        '  \"questions\": [\n'
        "    {\n"
        '      \"id\": number,\n'
        '      \"scores\": {\n'
        '        \"content\": number,\n'
        '        \"structure\": number,\n'
        '        \"clarity\": number,\n'
        '        \"confidence\": number,\n'
        '        \"relevance\": number\n'
        "      },\n"
        '      \"strengths\": [string, ...],\n'
        '      \"improvements\": [string, ...],\n'
        '      \"suggestedAnswer\": string\n'
        "    }\n"
        "  ],\n"
        '  \"overall\": {\n'
        '    \"overallScore\": number,\n'
        '    \"summary\": string,\n'
        '    \"strengths\": [string, ...],\n'
        '    \"improvements\": [string, ...]\n'
        "  }\n"
        "}\n\n"
        "No markdown. No explanations. No extra keys. JSON ONLY."
    )



@app.post("/api/score-interview", response_model=ScoreInterviewResp)
def score_interview(req: ScoreInterviewReq):
    if not req.answers:
        raise HTTPException(status_code=400, detail="No answers are provided to score.")

    try:
        system = build_scoring_system_prompt()

        user_payload = {
            "role": req.role,
            "difficulty": req.difficulty,
            "answers": [a.model_dump() for a in req.answers],
        }

        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
        ]

        r = call_openai_with_backoff(
            model="gpt-4o-mini",
            temperature=0.3,
            response_format={"type": "json_object"},
            messages=messages,
        )

        raw = r.choices[0].message.content
        data = json.loads(raw)

        q_list = data.get("questions") or []
        overall_raw = data.get("overall") or {}

        scored_by_id: Dict[str, Dict[str, Any]] = {}
        if isinstance(q_list, list):
            for item in q_list:
                try:
                    sid = str(item.get("id"))
                    if sid:
                        scored_by_id[sid] = item
                except Exception:
                    continue

        questions_out: List[QuestionFeedback] = []

        for a in req.answers:
            scored = scored_by_id.get(str(a.id), {}) or {}

            raw_scores = scored.get("scores") or {}

            def get_score(key: str, default: float = 3.0) -> float:
                try:
                    val = raw_scores.get(key, default)
                    return float(val)
                except Exception:
                    return default

            scores_obj = {
                "content":   get_score("content"),
                "structure": get_score("structure"),
                "clarity":   get_score("clarity"),
                "confidence": get_score("confidence"),
                "relevance": get_score("relevance"),
            }

            strengths = scored.get("strengths") or []
            if not isinstance(strengths, list):
                strengths = [str(strengths)]

            improvements = scored.get("improvements") or []
            if not isinstance(improvements, list):
                improvements = [str(improvements)]

            suggested = scored.get("suggestedAnswer") or ""
            if not suggested:
                suggested = (
                    a.idealAnswer
                    or "Ideal answer unavailable — The system took a coffee break mid-analysis."
                )

            if not strengths:
                strengths = ["Please try to answer the question, Build up your confidence."]
            if not improvements:
                improvements = [
                    "Think Netflix episode: start strong, show what you did, and end with a win — STAR style.",
                    "Throw in 1-2 numbers — %, time saved, $$$ — your answers glow instantly.",
                ]

            questions_out.append(
                QuestionFeedback(
                    id=a.id,
                    prompt=a.prompt,
                    interviewer=a.interviewer,
                    type=a.type,
                    userAnswer=a.userAnswer,
                    idealAnswer=a.idealAnswer
                    or "Ideal answer unavailable — The system took a coffee break mid-analysis.",
                    scores=scores_obj,
                    strengths=[str(x) for x in strengths],
                    improvements=[str(x) for x in improvements],
                    suggestedAnswer=suggested,
                )
            )

        try:
            overall_score = float(overall_raw.get("overallScore", 3.0))
        except Exception:
            overall_score = 3.0

        overall_summary = overall_raw.get("summary") or (
            "Overall summary unavailable — The system took a coffee break mid-analysis."
        )

        overall_strengths = overall_raw.get("strengths") or [
            "You stayed engaged and tried to answer every question.",
            "You've got the experience, shape it into punchy, high-impact stories",
        ]
        if not isinstance(overall_strengths, list):
            overall_strengths = [str(overall_strengths)]

        overall_improvements = overall_raw.get("improvements") or [
            "Lean on STAR — it's your cheat code for crisp, confident answers.",
            "Throw in 1-2 numbers — %, time saved, $$$ — your answers glow instantly.",
            "Take your time — a calm pause between points makes you sound like a pro.",
        ]
        if not isinstance(overall_improvements, list):
            overall_improvements = [str(overall_improvements)]

        overall_obj = OverallFeedback(
            overallScore=overall_score,
            summary=overall_summary,
            strengths=[str(x) for x in overall_strengths],
            improvements=[str(x) for x in overall_improvements],
        )

        # 🔹 This is your “report” object (I previously called it `report`)
        resp = ScoreInterviewResp(
            meta={
                "role": req.role,
                "difficulty": req.difficulty,
                "questionCount": len(req.answers),
                "fallback": False,
            },
            questions=questions_out,
            overall=overall_obj,
        )

        # 🔹 NEW: best-effort save to DB (does NOT break the API if DB fails)
        # 🔹 Save interview session to DB (best-effort, won’t break API if it fails)
        try:
            # derive interviewer names from the answers
            interviewer_names = sorted(
                {a.interviewer for a in req.answers if a.interviewer}
            )

            with Session(engine) as session:
                session.add(
                    InterviewSession(
                        role=req.role,
                        difficulty=req.difficulty,
                        question_count=len(req.answers),
                        interviewer_names=interviewer_names,
                        plan=req.plan,  # whatever the frontend sent us
                        answers=[a.model_dump() for a in req.answers],
                        report=resp.model_dump(),  # full scored report
                    )
                )
                session.commit()
        except Exception:
            logging.exception("Failed to persist interview session")


        return resp

    except Exception as e:
        logging.exception("score-interview failed")
        return ScoreInterviewResp(
            meta={
                "role": req.role,
                "difficulty": req.difficulty,
                "questionCount": len(req.answers),
                "fallback": True,
            },
            questions=[
                QuestionFeedback(
                    id=a.id,
                    prompt=a.prompt,
                    interviewer=a.interviewer,
                    type=a.type,
                    userAnswer=a.userAnswer,
                    idealAnswer=a.idealAnswer
                    or "Ideal answer unavailable — The system took a coffee break mid-analysis.",
                    scores={
                        "content": 3.0,
                        "structure": 3.0,
                        "clarity": 3.0,
                        "confidence": 3.0,
                        "relevance": 3.0,
                    },
                    strengths=[
                        "Please try to answer the question, Try to speak.",
                        "Say whatever you know, build up your confidence....",
                    ],
                    improvements=[
                        "Add more concrete examples and metrics.",
                        "Use a clear structure: situation, actions, and measurable result.",
                    ],
                    suggestedAnswer=(
                        "A good answer would briefly describe the situation, your actions, "
                        "and the measurable impact."
                    ),
                )
                for a in req.answers
            ],
            overall=OverallFeedback(
                overallScore=3.0,
                summary=(
                    "Promising baseline: Overall Feedback unavailable — "
                    "The system took a coffee break mid-analysis.."
                ),
                strengths=[
                    "You powered through every question — most users tap out early. Big respect for sticking with it!",
                    "Your resume shows real experience — now turn those moments into sharp, high-impact stories that land.",
                ],
                improvements=[
                    "Lean on STAR — it's your cheat code for crisp, confident answers.",
                    "Throw in 1-2 numbers — %, time saved, $$$ — your answers glow instantly.",
                    "Take your time — a calm pause between points makes you sound like a pro.",
                ],
            ),
        )


# -------------------- Interview Answer Endpoints --------------------
# Phase 1.1: Two-way communication - Answer Storage System

@app.post("/api/interview/answer/submit", response_model=SubmitAnswerResponse)
async def submit_interview_answer(request: SubmitAnswerRequest):
    """
    Submit and store a single interview answer.

    This endpoint:
    1. Validates that the referenced interview session exists
    2. Creates a new InterviewAnswer record in the database
    3. Returns the generated answer_id for reference

    Args:
        request: SubmitAnswerRequest containing all answer details

    Returns:
        SubmitAnswerResponse with success status and answer_id

    Raises:
        HTTPException 404: If the session_id doesn't exist
        HTTPException 500: If there's a database error
    """

    # Step 1: Validate that the interview session exists in the database
    with Session(engine) as db_session:
        existing_session = db_session.exec(
            select(InterviewSession).where(InterviewSession.id == request.session_id)
        ).first()

        # If no session found, return 404 error
        if not existing_session:
            raise HTTPException(
                status_code=404,
                detail=f"Interview session with id '{request.session_id}' not found"
            )

    # Step 2: Generate embedding for semantic search
    try:
        # Combine question and answer for better semantic context
        text_for_embedding = f"Question: {request.question_text}\nAnswer: {request.user_answer}"
        embedding = generate_embedding(text_for_embedding)
        embedding_json = json.dumps(embedding)
    except Exception as e:
        logging.warning(f"Failed to generate embedding: {str(e)}. Storing answer without embedding.")
        embedding_json = None

    # Step 3: Create and store the InterviewAnswer record
    try:
        with Session(engine) as db_session:
            # Create new answer object with all provided data
            new_answer = InterviewAnswer(
                session_id=request.session_id,
                question_id=request.question_id,
                question_text=request.question_text,
                question_intent=request.question_intent,
                role=request.role,
                user_answer=request.user_answer,
                transcript_raw=request.transcript_raw,
                audio_duration_seconds=request.audio_duration_seconds,
                embedding=embedding_json  # Store the embedding
            )

            # Add to database and commit the transaction
            db_session.add(new_answer)
            db_session.commit()

            # Refresh to get the auto-generated UUID
            db_session.refresh(new_answer)

            # Store the generated ID for the response
            answer_id = new_answer.id

    except Exception as e:
        # Log the error for debugging and return a generic 500 error
        logging.error(f"Failed to store interview answer: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to store interview answer. Please try again."
        )

    # Step 3: Return success response with the new answer's ID (convert UUID to string)
    return SubmitAnswerResponse(
        success=True,
        answer_id=str(answer_id),
        message="Answer submitted successfully"
    )


@app.get("/api/interview/answers/{session_id}", response_model=RetrieveAnswersResponse)
async def get_interview_answers(session_id: str):
    """
    Retrieve all answers for a given interview session.

    This endpoint:
    1. Validates that the interview session exists
    2. Fetches all answers associated with the session
    3. Returns them ordered by answer_timestamp (chronological)

    Args:
        session_id: UUID of the interview session

    Returns:
        RetrieveAnswersResponse with list of all answers

    Raises:
        HTTPException 404: If the session_id doesn't exist
        HTTPException 500: If there's a database error
    """

    # Step 1: Validate that the interview session exists
    with Session(engine) as db_session:
        existing_session = db_session.exec(
            select(InterviewSession).where(InterviewSession.id == session_id)
        ).first()

        # If no session found, return 404 error
        if not existing_session:
            raise HTTPException(
                status_code=404,
                detail=f"Interview session with id '{session_id}' not found"
            )

    # Step 2: Retrieve all answers for this session, ordered chronologically
    try:
        with Session(engine) as db_session:
            # Query answers filtered by session_id, ordered by answer_timestamp
            answers = db_session.exec(
                select(InterviewAnswer)
                .where(InterviewAnswer.session_id == session_id)
                .order_by(InterviewAnswer.answer_timestamp)
            ).all()

            # Convert SQLModel objects to Pydantic response models (convert UUIDs to strings)
            answer_list = [
                AnswerDetail(
                    id=str(answer.id),
                    session_id=str(answer.session_id),
                    question_id=answer.question_id,
                    question_text=answer.question_text,
                    question_intent=answer.question_intent,
                    role=answer.role,
                    user_answer=answer.user_answer,
                    transcript_raw=answer.transcript_raw,
                    audio_duration_seconds=float(answer.audio_duration_seconds) if answer.audio_duration_seconds else None,
                    answer_timestamp=answer.answer_timestamp,
                    created_at=answer.created_at
                )
                for answer in answers
            ]

    except Exception as e:
        # Log the error and return a 500 response
        logging.error(f"Failed to retrieve interview answers: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to retrieve interview answers. Please try again."
        )

    # Step 3: Return the list of answers
    return RetrieveAnswersResponse(
        success=True,
        session_id=session_id,
        total_answers=len(answer_list),
        answers=answer_list
    )


@app.post("/api/interview/session/create", response_model=CreateSessionResponse)
async def create_interview_session(request: CreateSessionRequest):
    """
    Create a new interview session at the START of an interview.

    This endpoint:
    1. Creates a new InterviewSession record with the provided metadata
    2. Returns the session_id for use in subsequent answer submissions

    This allows answers to be saved in real-time during the interview,
    rather than waiting until the end. If the browser crashes, answers
    are already persisted in the database.

    Args:
        request: CreateSessionRequest with interview configuration

    Returns:
        CreateSessionResponse with the new session_id

    Raises:
        HTTPException 500: If there's a database error
    """

    try:
        with Session(engine) as db_session:
            # Create new interview session with provided metadata
            new_session = InterviewSession(
                role=request.role,
                difficulty=request.difficulty,
                question_count=request.question_count,
                interviewer_names=request.interviewer_names or [],
                plan=request.plan,
                # answers and report will be updated later
                answers=[],
                report=None
            )

            # Add to database and commit
            db_session.add(new_session)
            db_session.commit()

            # Refresh to get the auto-generated UUID
            db_session.refresh(new_session)

            session_id = new_session.id

    except Exception as e:
        # Log the error and return a 500 response
        logging.error(f"Failed to create interview session: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to create interview session. Please try again."
        )

    # Return success with the new session_id (convert UUID to string)
    return CreateSessionResponse(
        success=True,
        session_id=str(session_id),
        message="Interview session created successfully"
    )


@app.get("/api/interview/session/{session_id}", response_model=SessionDetailResponse)
async def get_interview_session(session_id: str):
    """
    Retrieve full session details including all answers.

    This endpoint is used by the Feedback page to fetch interview data
    from the database instead of relying on localStorage.

    Args:
        session_id: UUID of the interview session

    Returns:
        SessionDetailResponse with session metadata, plan, and all answers

    Raises:
        HTTPException 404: If the session_id doesn't exist
        HTTPException 500: If there's a database error
    """

    try:
        with Session(engine) as db_session:
            # Fetch the interview session
            interview_session = db_session.exec(
                select(InterviewSession).where(InterviewSession.id == session_id)
            ).first()

            # If no session found, return 404 error
            if not interview_session:
                raise HTTPException(
                    status_code=404,
                    detail=f"Interview session with id '{session_id}' not found"
                )

            # Fetch all answers for this session, ordered chronologically
            answers = db_session.exec(
                select(InterviewAnswer)
                .where(InterviewAnswer.session_id == session_id)
                .order_by(InterviewAnswer.answer_timestamp)
            ).all()

            # Convert answers to response format (convert UUIDs to strings)
            answer_list = [
                AnswerDetail(
                    id=str(answer.id),
                    session_id=str(answer.session_id),
                    question_id=answer.question_id,
                    question_text=answer.question_text,
                    question_intent=answer.question_intent,
                    role=answer.role,
                    user_answer=answer.user_answer,
                    transcript_raw=answer.transcript_raw,
                    audio_duration_seconds=float(answer.audio_duration_seconds) if answer.audio_duration_seconds else None,
                    answer_timestamp=answer.answer_timestamp,
                    created_at=answer.created_at
                )
                for answer in answers
            ]

            # Return full session details (convert UUID to string)
            return SessionDetailResponse(
                success=True,
                session_id=str(interview_session.id),
                role=interview_session.role,
                difficulty=interview_session.difficulty,
                question_count=interview_session.question_count,
                interviewer_names=interview_session.interviewer_names or [],
                plan=interview_session.plan,
                answers=answer_list,
                created_at=interview_session.created_at
            )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Failed to retrieve interview session: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to retrieve interview session. Please try again."
        )


# -------------------- Semantic Search Endpoint --------------------

@app.post("/api/interview/search-answers", response_model=SearchAnswersResponse)
async def search_interview_answers(request: SearchAnswersRequest):
    """
    Semantic search for similar answers within an interview session.

    This endpoint:
    1. Validates that the referenced interview session exists
    2. Generates an embedding for the search query
    3. Finds the top_k most similar answers using cosine similarity
    4. Returns results sorted by similarity score (highest first)

    Use cases:
    - Find answers related to a specific topic (e.g., "machine learning")
    - Identify patterns in how candidate answered similar questions
    - Enable follow-up question generation based on previous answers

    Args:
        request: SearchAnswersRequest with session_id, query, and top_k

    Returns:
        SearchAnswersResponse with ranked list of similar answers

    Raises:
        HTTPException 404: If the session_id doesn't exist
        HTTPException 400: If the query is empty
        HTTPException 500: If embedding generation or search fails
    """

    # Step 1: Validate that the query is not empty
    if not request.query or not request.query.strip():
        raise HTTPException(
            status_code=400,
            detail="Search query cannot be empty"
        )

    # Step 2: Validate that the interview session exists in the database
    with Session(engine) as db_session:
        existing_session = db_session.exec(
            select(InterviewSession).where(InterviewSession.id == request.session_id)
        ).first()

        if not existing_session:
            raise HTTPException(
                status_code=404,
                detail=f"Interview session with id '{request.session_id}' not found"
            )

    # Step 3: Generate embedding for the search query
    try:
        query_embedding = generate_embedding(request.query)
    except Exception as e:
        logging.error(f"Failed to generate query embedding: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to process search query. Please try again."
        )

    # Step 4: Find similar answers using embedding similarity
    try:
        similar_answers = find_similar_answers(
            session_id=request.session_id,
            query_embedding=query_embedding,
            top_k=request.top_k
        )
    except Exception as e:
        logging.error(f"Failed to search answers: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Failed to search answers. Please try again."
        )

    # Step 5: Transform results to response format
    results = [
        SearchAnswerResult(
            answer_id=answer["answer_id"],
            question_id=answer["question_id"],
            question_text=answer["question_text"],
            user_answer=answer["user_answer"],
            role=answer["role"],
            similarity_score=answer["similarity"]
        )
        for answer in similar_answers
    ]

    # Step 6: Return the search results
    return SearchAnswersResponse(
        success=True,
        session_id=request.session_id,
        query=request.query,
        total_results=len(results),
        results=results
    )


# -------------------- Phase 1.3: Intelligent Interview Flow Endpoints --------------------
# These endpoints provide intelligent, context-aware question generation

# Global orchestrator instance (lazy initialization)
_orchestrator_instance: Optional[InterviewOrchestrator] = None


def get_orchestrator() -> InterviewOrchestrator:
    """Get/create the InterviewOrchestrator instance with lazy initialization."""
    global _orchestrator_instance
    if _orchestrator_instance is None:
        logging.info("Initializing InterviewOrchestrator...")
        _orchestrator_instance = InterviewOrchestrator()
        logging.info("InterviewOrchestrator initialized successfully")
    return _orchestrator_instance


def validate_difficulty(difficulty: str) -> str:
    """Validate and normalize difficulty level."""
    valid_difficulties = ["easy", "medium", "hard", "junior", "intermediate", "senior", "intern", "associate", "lead"]
    difficulty_lower = difficulty.lower().strip()
    if difficulty_lower not in valid_difficulties:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid difficulty '{difficulty}'. Must be one of: {', '.join(valid_difficulties)}"
        )
    return difficulty_lower


def validate_uuid(uuid_string: str, field_name: str = "session_id") -> str:
    """Validate UUID format."""
    import uuid as uuid_module
    try:
        uuid_module.UUID(uuid_string)
        return uuid_string
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid {field_name} format. Must be a valid UUID."
        )


@app.post("/api/interview/next-question", response_model=NextQuestionResponse)
async def get_next_question(request: NextQuestionRequest):
    """
    Get next contextually-aware interview question.

    Uses the InterviewOrchestrator to generate intelligent questions based on
    conversation history, detected patterns, and semantic analysis.
    """
    logging.info(f"POST /api/interview/next-question - Session: {request.session_id}, Q#{request.current_question_number}")

    try:
        # Validate inputs
        session_id = validate_uuid(request.session_id, "session_id")
        difficulty = validate_difficulty(request.difficulty)

        # Verify session exists
        with Session(engine) as db_session:
            existing_session = db_session.exec(
                select(InterviewSession).where(InterviewSession.id == session_id)
            ).first()
            if not existing_session:
                raise HTTPException(status_code=404, detail=f"Interview session '{session_id}' not found")

        # Get orchestrator and generate question
        orchestrator = get_orchestrator()
        from uuid import UUID as UUIDType

        result = await orchestrator.get_next_question(
            session_id=UUIDType(session_id),
            current_question_number=request.current_question_number,
            role=request.role,
            difficulty=difficulty,
            total_questions=request.total_questions
        )

        logging.info(f"Generated question type: {result.get('question', {}).get('type', 'unknown')}")

        # Transform result to response model
        question_data = result.get("question", {})
        references_data = result.get("references", {})
        metadata_data = result.get("metadata", {})

        return NextQuestionResponse(
            question=QuestionData(
                id=question_data.get("id", request.current_question_number),
                text=question_data.get("text", ""),
                intent=question_data.get("intent", "general"),
                type=question_data.get("type", "standard")
            ),
            interviewer_comment=result.get("interviewer_comment"),
            references=QuestionReferences(
                question_id=references_data.get("question_id"),
                excerpt=references_data.get("excerpt")
            ),
            metadata=QuestionMetadata(
                decision_reason=metadata_data.get("decision_reason", ""),
                patterns_detected=metadata_data.get("patterns_detected", []),
                conversation_stage=metadata_data.get("conversation_stage", "unknown"),
                action_taken=metadata_data.get("action_taken", "standard"),
                context_used=metadata_data.get("context_used"),
                generated_at=metadata_data.get("generated_at")
            )
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error generating next question: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate next question: {str(e)}")


@app.post("/api/interview/submit-and-next", response_model=SubmitAndNextResponse)
async def submit_answer_and_get_next(request: SubmitAndNextRequest):
    """
    Submit an answer AND get the next question in a single call.
    Provides seamless conversation flow.
    """
    logging.info(f"POST /api/interview/submit-and-next - Session: {request.session_id}, Q#{request.question_id}")

    answer_stored = False
    answer_id = None
    embedding_generated = False

    try:
        # Validate inputs
        session_id = validate_uuid(request.session_id, "session_id")
        difficulty = validate_difficulty(request.difficulty)

        # Verify session exists
        with Session(engine) as db_session:
            existing_session = db_session.exec(
                select(InterviewSession).where(InterviewSession.id == session_id)
            ).first()
            if not existing_session:
                raise HTTPException(status_code=404, detail=f"Interview session '{session_id}' not found")

        # Store the answer (best effort)
        try:
            text_for_embedding = f"Question: {request.question_text}\nAnswer: {request.user_answer}"
            try:
                embedding = generate_embedding(text_for_embedding)
                embedding_json = json.dumps(embedding)
                embedding_generated = True
            except Exception as embed_error:
                logging.warning(f"Failed to generate embedding: {embed_error}")
                embedding_json = None

            with Session(engine) as db_session:
                new_answer = InterviewAnswer(
                    session_id=session_id,
                    question_id=request.question_id,
                    question_text=request.question_text,
                    question_intent=request.question_intent,
                    role=request.role,
                    user_answer=request.user_answer,
                    transcript_raw=request.transcript_raw,
                    audio_duration_seconds=request.audio_duration_seconds,
                    embedding=embedding_json
                )
                db_session.add(new_answer)
                db_session.commit()
                db_session.refresh(new_answer)
                answer_id = str(new_answer.id)
                answer_stored = True
                logging.info(f"Answer stored with ID: {answer_id}")
        except Exception as store_error:
            logging.error(f"Failed to store answer: {store_error}")

        # Generate next question
        orchestrator = get_orchestrator()
        next_question_number = request.question_id + 1
        from uuid import UUID as UUIDType

        result = await orchestrator.get_next_question(
            session_id=UUIDType(session_id),
            current_question_number=next_question_number,
            role=request.role,
            difficulty=difficulty,
            total_questions=request.total_questions
        )

        question_data = result.get("question", {})
        references_data = result.get("references", {})
        metadata_data = result.get("metadata", {})

        next_question = NextQuestionResponse(
            question=QuestionData(
                id=question_data.get("id", next_question_number),
                text=question_data.get("text", ""),
                intent=question_data.get("intent", "general"),
                type=question_data.get("type", "standard")
            ),
            interviewer_comment=result.get("interviewer_comment"),
            references=QuestionReferences(
                question_id=references_data.get("question_id"),
                excerpt=references_data.get("excerpt")
            ),
            metadata=QuestionMetadata(
                decision_reason=metadata_data.get("decision_reason", ""),
                patterns_detected=metadata_data.get("patterns_detected", []),
                conversation_stage=metadata_data.get("conversation_stage", "unknown"),
                action_taken=metadata_data.get("action_taken", "standard"),
                context_used=metadata_data.get("context_used"),
                generated_at=metadata_data.get("generated_at")
            )
        )

        return SubmitAndNextResponse(
            answer_stored=answer_stored,
            answer_id=answer_id,
            embedding_generated=embedding_generated,
            next_question=next_question
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error in submit-and-next: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process: {str(e)}")


@app.get("/api/interview/conversation-state/{session_id}", response_model=ConversationStateResponse)
async def get_conversation_state(session_id: str):
    """
    Get current state of interview conversation.
    Useful for debugging and monitoring interview quality.
    """
    logging.info(f"GET /api/interview/conversation-state/{session_id}")

    try:
        session_id = validate_uuid(session_id, "session_id")

        # Verify session exists
        with Session(engine) as db_session:
            existing_session = db_session.exec(
                select(InterviewSession).where(InterviewSession.id == session_id)
            ).first()
            if not existing_session:
                raise HTTPException(status_code=404, detail=f"Interview session '{session_id}' not found")

        # Analyze conversation
        orchestrator = get_orchestrator()
        from uuid import UUID as UUIDType

        result = await orchestrator.analyze_conversation_state(session_id=UUIDType(session_id))
        logging.info(f"Conversation state analyzed: {result.get('total_answers', 0)} answers")

        # Transform contradictions
        contradictions = [
            ContradictionDetail(
                previous_statement=c.get("previous_statement"),
                current_statement=c.get("current_statement"),
                contradiction_type=c.get("contradiction_type"),
                confidence_score=c.get("confidence_score"),
                explanation=c.get("explanation")
            )
            for c in result.get("contradictions_detected", [])
        ]

        raw_metrics = result.get("quality_metrics", {})
        quality_metrics = QualityMetrics(
            avg_answer_length=raw_metrics.get("avg_answer_length", 0),
            star_format_usage=raw_metrics.get("star_format_usage", 0),
            specificity_score=raw_metrics.get("specificity_score", 0)
        )

        return ConversationStateResponse(
            total_answers=result.get("total_answers", 0),
            topics_discussed=result.get("topics_discussed", []),
            repeated_topics=result.get("repeated_topics", {}),
            contradictions_detected=contradictions,
            conversation_summary=result.get("conversation_summary", ""),
            quality_metrics=quality_metrics,
            recommendations=result.get("recommendations", [])
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error analyzing conversation state: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze conversation: {str(e)}")


@app.get("/api/interview/patterns/{session_id}", response_model=PatternsResponse)
async def get_conversation_patterns(session_id: str):
    """
    Get detected patterns in conversation.
    Quick endpoint for pattern analysis.
    """
    logging.info(f"GET /api/interview/patterns/{session_id}")

    try:
        session_id = validate_uuid(session_id, "session_id")

        # Verify session exists
        with Session(engine) as db_session:
            existing_session = db_session.exec(
                select(InterviewSession).where(InterviewSession.id == session_id)
            ).first()
            if not existing_session:
                raise HTTPException(status_code=404, detail=f"Interview session '{session_id}' not found")

        # Get all answers
        all_answers = get_all_answers(session_id)
        if not all_answers:
            return PatternsResponse(
                contradictions=[],
                repeated_topics={},
                weak_answers=[],
                strong_areas=[],
                gaps=[]
            )

        # Detect repeated topics
        repeated_topics = detect_repeated_topics(session_id)

        # Detect contradictions
        contradictions = []
        if len(all_answers) >= 2:
            last_answer = all_answers[-1]
            try:
                raw_contradictions = await detect_contradictions(
                    session_id,
                    last_answer.get("user_answer", ""),
                    last_answer.get("question_text", "")
                )
                for c in raw_contradictions:
                    contradictions.append(ContradictionDetail(
                        previous_statement=c.get("previous_statement"),
                        current_statement=c.get("current_statement"),
                        contradiction_type=c.get("contradiction_type"),
                        confidence_score=c.get("confidence_score"),
                        explanation=c.get("explanation")
                    ))
            except Exception as e:
                logging.warning(f"Failed to detect contradictions: {e}")

        # Analyze answer quality
        weak_answers = []
        orchestrator = get_orchestrator()
        for answer in all_answers:
            quality = orchestrator.decision_engine._analyze_answer_quality(
                answer.get("user_answer", ""),
                answer.get("question_intent", "general")
            )
            if quality.get("completeness_score", 1) < 0.4 or quality.get("is_vague", False):
                weak_answers.append(answer.get("question_id", 0))

        # Identify strong areas
        strong_areas = [topic for topic, count in repeated_topics.items() if count >= 2]

        # Identify gaps
        expected_topics = ["experience", "skills", "challenges", "achievements", "teamwork"]
        topics_discussed = extract_topics(session_id)
        topics_lower = [t.lower() for t in topics_discussed]
        gaps = [exp for exp in expected_topics if not any(exp in topic for topic in topics_lower)]

        return PatternsResponse(
            contradictions=contradictions,
            repeated_topics=repeated_topics,
            weak_answers=weak_answers,
            strong_areas=strong_areas,
            gaps=gaps
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error detecting patterns: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to detect patterns: {str(e)}")


# ==================== Phase 1.3: TTS (Text-to-Speech) Endpoints ====================

# TTS Service singleton
_tts_service: Optional[TTSService] = None

def get_tts() -> TTSService:
    """Get or create TTS service singleton."""
    global _tts_service
    if _tts_service is None:
        _tts_service = get_tts_service()
    return _tts_service


class TTSGenerateRequest(BaseModel):
    """Request model for TTS generation."""
    text: str = Field(..., description="Text to convert to speech", min_length=1, max_length=4096)
    context: str = Field(
        default="question",
        description="Context type: question, acknowledgment, follow_up, challenge, encouragement"
    )
    voice: Optional[str] = Field(
        default=None,
        description="Voice: alloy, echo, fable, onyx, nova, shimmer"
    )
    speed: Optional[float] = Field(
        default=None,
        ge=0.25,
        le=4.0,
        description="Speech speed (0.25-4.0)"
    )
    conversation_stage: str = Field(
        default="mid",
        description="Interview stage: early, mid, late"
    )
    return_url: bool = Field(
        default=True,
        description="If true, return URL to cached file. If false, return base64 audio."
    )


class TTSGenerateResponse(BaseModel):
    """Response model for TTS generation."""
    success: bool
    audio_url: Optional[str] = None
    audio_base64: Optional[str] = None
    filename: Optional[str] = None
    duration_estimate_ms: Optional[int] = None
    voice_used: str
    cached: bool = False
    error: Optional[str] = None


@app.post("/api/tts/generate", response_model=TTSGenerateResponse)
async def generate_tts(request: TTSGenerateRequest):
    """
    Generate Text-to-Speech audio from text.

    Converts interviewer text to natural speech using OpenAI TTS API.
    Supports different voices, speeds, and interview contexts.

    Request body:
    - text: The text to convert (required, 1-4096 chars)
    - context: Type of content (question, acknowledgment, follow_up, etc.)
    - voice: Voice selection (alloy, echo, fable, onyx, nova, shimmer)
    - speed: Speech rate (0.25 to 4.0)
    - conversation_stage: Interview stage (early, mid, late)
    - return_url: If true, return URL; if false, return base64

    Returns:
    - audio_url: URL to fetch the audio file (if return_url=true)
    - audio_base64: Base64-encoded audio (if return_url=false)
    - filename: Name of the cached file
    - voice_used: Which voice was used
    - cached: Whether result was from cache
    """
    import base64

    logging.info(f"TTS request: {len(request.text)} chars, context={request.context}")

    try:
        tts = get_tts()

        # Check if should speak this
        if not tts.should_speak_this(request.text, request.context):
            logging.info("TTS skipped: should_not_speak")
            return TTSGenerateResponse(
                success=False,
                error="Text should not be spoken (too long, system message, or code)",
                voice_used="none",
                cached=False
            )

        # Generate cache key
        cache_key = f"{request.context}_{hash(request.text) % 100000}"

        # Build kwargs for generation
        kwargs = {}
        if request.voice:
            kwargs["voice"] = request.voice
        if request.speed:
            kwargs["speed"] = request.speed

        # Try to use cached version first
        try:
            audio_bytes, file_path = await tts.generate_and_cache(
                request.text,
                cache_key,
                **kwargs
            )
            cached = True
        except Exception:
            # Fallback to direct generation with context
            audio_bytes = await tts.generate_for_interview_context(
                request.text,
                request.context,
                request.conversation_stage
            )
            # Save to file
            file_path = await tts.save_audio_file(
                audio_bytes,
                cache_key
            )
            cached = False

        # Get filename from path
        filename = Path(file_path).name

        # Determine voice used
        voice_used = request.voice or tts._select_voice_for_context(request.context)

        # Estimate duration (rough: ~150 words per minute, ~5 chars per word)
        word_count = len(request.text.split())
        duration_estimate_ms = int((word_count / 150) * 60 * 1000)

        if request.return_url:
            # Return URL to audio file
            audio_url = f"/api/audio/{filename}"
            return TTSGenerateResponse(
                success=True,
                audio_url=audio_url,
                filename=filename,
                duration_estimate_ms=duration_estimate_ms,
                voice_used=voice_used,
                cached=cached
            )
        else:
            # Return base64-encoded audio
            audio_base64 = base64.b64encode(audio_bytes).decode("utf-8")
            return TTSGenerateResponse(
                success=True,
                audio_base64=audio_base64,
                filename=filename,
                duration_estimate_ms=duration_estimate_ms,
                voice_used=voice_used,
                cached=cached
            )

    except Exception as e:
        logging.exception(f"TTS generation failed: {str(e)}")
        return TTSGenerateResponse(
            success=False,
            error=str(e),
            voice_used="none",
            cached=False
        )


@app.get("/api/audio/{filename}")
async def serve_audio(filename: str):
    """
    Serve cached audio files.

    Returns the audio file for playback in the browser.
    Supports MP3 format with proper MIME type.

    Path params:
    - filename: Name of the audio file (e.g., "question_12345_abc123.mp3")

    Returns:
    - Audio file with Content-Type: audio/mpeg
    """
    # Validate filename (security: prevent path traversal)
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    # Only allow .mp3 files
    if not filename.endswith(".mp3"):
        raise HTTPException(status_code=400, detail="Only MP3 files are supported")

    # Build file path
    audio_dir = Path("audio_cache")
    file_path = audio_dir / filename

    # Check if file exists
    if not file_path.exists():
        logging.warning(f"Audio file not found: {filename}")
        raise HTTPException(status_code=404, detail="Audio file not found")

    # Serve the file
    return FileResponse(
        path=str(file_path),
        media_type="audio/mpeg",
        filename=filename,
        headers={
            "Cache-Control": "public, max-age=3600",  # Cache for 1 hour
            "Accept-Ranges": "bytes"  # Support partial content for seeking
        }
    )


class TTSBatchRequest(BaseModel):
    """Request model for batch TTS generation."""
    items: List[Dict[str, str]] = Field(
        ...,
        description="List of items: [{text, context, id}, ...]"
    )
    max_concurrent: int = Field(default=3, ge=1, le=5)


class TTSBatchItemResult(BaseModel):
    """Result for a single batch item."""
    id: str
    success: bool
    audio_url: Optional[str] = None
    error: Optional[str] = None


class TTSBatchResponse(BaseModel):
    """Response model for batch TTS generation."""
    success: bool
    results: List[TTSBatchItemResult]
    total: int
    successful: int
    failed: int


@app.post("/api/tts/batch", response_model=TTSBatchResponse)
async def generate_tts_batch(request: TTSBatchRequest):
    """
    Generate TTS for multiple texts in batch.

    Useful for pre-generating audio for multiple questions.

    Request body:
    - items: List of {text, context, id} objects
    - max_concurrent: Max concurrent API calls (1-5)

    Returns:
    - results: List of results with audio URLs
    - total/successful/failed counts
    """
    logging.info(f"TTS batch request: {len(request.items)} items")

    try:
        tts = get_tts()

        # Generate batch
        raw_results = await tts.generate_batch(
            request.items,
            max_concurrent=request.max_concurrent
        )

        # Format results
        results = []
        successful = 0
        failed = 0

        for result in raw_results:
            item_id = result.get("id", "unknown")

            if result.get("success"):
                file_path = result.get("file_path", "")
                filename = Path(file_path).name if file_path else None
                audio_url = f"/api/audio/{filename}" if filename else None

                results.append(TTSBatchItemResult(
                    id=item_id,
                    success=True,
                    audio_url=audio_url
                ))
                successful += 1
            else:
                results.append(TTSBatchItemResult(
                    id=item_id,
                    success=False,
                    error=result.get("error", result.get("reason", "Unknown error"))
                ))
                failed += 1

        return TTSBatchResponse(
            success=failed == 0,
            results=results,
            total=len(request.items),
            successful=successful,
            failed=failed
        )

    except Exception as e:
        logging.exception(f"TTS batch generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Batch TTS failed: {str(e)}")


# ==================== Phase 1.3: Real-time Conversational Endpoints ====================
# These endpoints enable full voice conversation experience with TTS support

# ---------- Pydantic Models for Real-time Conversation ----------

class SubmitAnswerRealtimeRequest(BaseModel):
    """Request model for submitting answer with real-time AI response."""
    session_id: str = Field(..., description="Interview session UUID")
    question_id: int = Field(..., ge=1, description="Current question number")
    question_text: str = Field(..., min_length=1, description="The question that was asked")
    question_intent: str = Field(default="behavioral", description="Question intent type")
    role: str = Field(..., min_length=1, description="Job role being interviewed for")
    user_answer: str = Field(..., min_length=1, description="User's answer text")
    transcript_raw: Optional[str] = Field(default=None, description="Raw transcript from STT")
    audio_duration_seconds: Optional[float] = Field(default=None, ge=0, description="Duration of audio answer")
    difficulty: str = Field(default="medium", description="Interview difficulty level")
    total_questions: int = Field(default=10, ge=1, le=50, description="Total planned questions")
    generate_audio: bool = Field(default=True, description="Whether to generate TTS audio")


class AcknowledgmentModel(BaseModel):
    """Model for AI acknowledgment response."""
    text: str
    audio_url: Optional[str] = None
    should_speak: bool = True
    tone: str = "neutral"


class FollowUpProbeModel(BaseModel):
    """Model for follow-up probe."""
    text: str
    audio_url: Optional[str] = None
    probe_type: str = "specific"
    missing_element: Optional[str] = None


class TransitionModel(BaseModel):
    """Model for transition between questions."""
    text: str
    audio_url: Optional[str] = None


class AIResponseModel(BaseModel):
    """Model for complete AI response."""
    acknowledgment: Optional[AcknowledgmentModel] = None
    follow_up_probe: Optional[FollowUpProbeModel] = None
    transition: Optional[TransitionModel] = None


class FlowControlModel(BaseModel):
    """Model for conversation flow control."""
    should_proceed_to_next: bool
    needs_follow_up: bool
    quality_sufficient: bool
    conversation_stage: str


class QuestionWithAudioModel(BaseModel):
    """Model for question with optional audio."""
    id: int
    text: str
    intent: str
    type: str
    audio_url: Optional[str] = None


class NextQuestionWithAudioModel(BaseModel):
    """Model for next question response with audio."""
    question: QuestionWithAudioModel
    interviewer_comment: Optional[str] = None
    interviewer_comment_audio_url: Optional[str] = None
    references: Optional[Dict[str, Any]] = None
    metadata: Optional[Dict[str, Any]] = None


class SubmitAnswerRealtimeResponse(BaseModel):
    """Response model for real-time answer submission."""
    answer_stored: bool
    answer_id: Optional[str] = None
    ai_response: AIResponseModel
    next_question: Optional[NextQuestionWithAudioModel] = None
    flow_control: FlowControlModel
    quality_metrics: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class SubmitFollowUpRequest(BaseModel):
    """Request model for submitting follow-up answer."""
    session_id: str = Field(..., description="Interview session UUID")
    original_question_id: int = Field(..., ge=1, description="Original question ID")
    original_question_text: str = Field(..., description="Original question text")
    original_question_intent: str = Field(default="behavioral", description="Original question intent")
    follow_up_answer: str = Field(..., min_length=1, description="User's follow-up elaboration")
    role: str = Field(..., description="Job role")
    difficulty: str = Field(default="medium", description="Difficulty level")
    total_questions: int = Field(default=10, description="Total questions")
    generate_audio: bool = Field(default=True, description="Generate TTS audio")


class SubmitFollowUpResponse(BaseModel):
    """Response model for follow-up submission."""
    follow_up_stored: bool
    answer_id: Optional[str] = None
    is_follow_up_response: bool = True
    original_question_id: int
    follow_up_attempt: int
    ai_response: AIResponseModel
    next_question: Optional[NextQuestionWithAudioModel] = None
    flow_control: FlowControlModel
    quality_metrics: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class StartInterviewWithAudioRequest(BaseModel):
    """Request model for starting interview with audio."""
    session_id: str = Field(..., description="Interview session UUID")
    role: str = Field(..., min_length=1, description="Job role")
    difficulty: str = Field(default="medium", description="Difficulty level")
    total_questions: int = Field(default=10, ge=1, le=50, description="Total questions")
    generate_audio: bool = Field(default=True, description="Generate TTS audio")


class StartInterviewWithAudioResponse(BaseModel):
    """Response model for interview start with audio."""
    interview_started: bool
    session_id: str
    first_question: Optional[NextQuestionWithAudioModel] = None
    error: Optional[str] = None


class RegenerateAudioRequest(BaseModel):
    """Request model for regenerating audio."""
    text: str = Field(..., min_length=1, max_length=4096, description="Text to convert")
    context_type: str = Field(default="question", description="Context type for TTS")
    voice: Optional[str] = Field(default=None, description="Voice selection")
    conversation_stage: str = Field(default="mid", description="Conversation stage")
    force_regenerate: bool = Field(default=False, description="Force regeneration even if cached")


class RegenerateAudioResponse(BaseModel):
    """Response model for audio regeneration."""
    audio_generated: bool
    audio_url: Optional[str] = None
    cached: bool = False
    error: Optional[str] = None


class AudioStatusResponse(BaseModel):
    """Response model for audio status."""
    session_id: str
    cache_stats: Dict[str, Any]


class ClearCacheRequest(BaseModel):
    """Request model for clearing cache."""
    older_than_days: int = Field(default=7, ge=1, le=365, description="Delete files older than X days")


class ClearCacheResponse(BaseModel):
    """Response model for cache clearing."""
    files_deleted: int
    space_freed_mb: float


# ---------- Real-time Conversational Endpoints ----------

@app.post("/api/interview/submit-answer-realtime", response_model=SubmitAnswerRealtimeResponse)
async def submit_answer_realtime(request: SubmitAnswerRealtimeRequest):
    """
    Submit answer and get immediate AI response with audio.

    This is the MAIN endpoint for conversational interviews with TTS support.
    Processes the answer, analyzes quality, generates AI response with audio,
    and returns the next question (if ready to proceed).

    Flow:
    1. Validate inputs
    2. Process answer through orchestrator
    3. Generate AI response (acknowledgment, probe, transition)
    4. Convert responses to audio (if generate_audio=True)
    5. Generate next question with audio (if should_proceed)
    6. Return comprehensive response
    """
    start_time = time.time()
    logging.info(f"POST /api/interview/submit-answer-realtime - Session: {request.session_id}, Q{request.question_id}")

    try:
        # Validate session_id
        validate_uuid(request.session_id, "session_id")

        # Validate difficulty
        difficulty = validate_difficulty(request.difficulty)

        # Get orchestrator
        orchestrator = get_orchestrator()

        # Build answer data
        answer_data = {
            "question_id": request.question_id,
            "question_text": request.question_text,
            "question_intent": request.question_intent,
            "user_answer": request.user_answer,
            "transcript_raw": request.transcript_raw or request.user_answer,
            "audio_duration_seconds": request.audio_duration_seconds or 0
        }

        # Process answer with real-time response
        result = await orchestrator.process_answer_with_realtime_response(
            session_id=UUID(request.session_id),
            answer_data=answer_data,
            role=request.role,
            difficulty=difficulty,
            total_questions=request.total_questions,
            generate_audio=request.generate_audio
        )

        # Build response models
        ai_response_data = result.get("ai_response", {})
        ai_response = AIResponseModel(
            acknowledgment=AcknowledgmentModel(**ai_response_data["acknowledgment"]) if ai_response_data.get("acknowledgment") else None,
            follow_up_probe=FollowUpProbeModel(**ai_response_data["follow_up_probe"]) if ai_response_data.get("follow_up_probe") else None,
            transition=TransitionModel(**ai_response_data["transition"]) if ai_response_data.get("transition") else None
        )

        flow_control_data = result.get("flow_control", {})
        flow_control = FlowControlModel(
            should_proceed_to_next=flow_control_data.get("should_proceed_to_next", True),
            needs_follow_up=flow_control_data.get("needs_follow_up", False),
            quality_sufficient=flow_control_data.get("quality_sufficient", True),
            conversation_stage=flow_control_data.get("conversation_stage", "mid")
        )

        # Build next question model if present
        next_question = None
        if result.get("next_question"):
            nq = result["next_question"]
            q = nq.get("question", {})
            next_question = NextQuestionWithAudioModel(
                question=QuestionWithAudioModel(
                    id=q.get("id", request.question_id + 1),
                    text=q.get("text", ""),
                    intent=q.get("intent", "general"),
                    type=q.get("type", "standard"),
                    audio_url=q.get("audio_url")
                ),
                interviewer_comment=nq.get("interviewer_comment"),
                interviewer_comment_audio_url=nq.get("interviewer_comment_audio_url"),
                references=nq.get("references"),
                metadata=nq.get("metadata")
            )

        elapsed = time.time() - start_time
        logging.info(f"Submit-answer-realtime completed in {elapsed:.2f}s - Proceed: {flow_control.should_proceed_to_next}")

        return SubmitAnswerRealtimeResponse(
            answer_stored=result.get("answer_stored", True),
            answer_id=result.get("answer_id"),
            ai_response=ai_response,
            next_question=next_question,
            flow_control=flow_control,
            quality_metrics=result.get("quality_metrics"),
            error=result.get("error")
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error in submit-answer-realtime: {str(e)}")
        # Return graceful error response
        return SubmitAnswerRealtimeResponse(
            answer_stored=False,
            answer_id=None,
            ai_response=AIResponseModel(
                acknowledgment=AcknowledgmentModel(
                    text="Thank you for your response.",
                    audio_url=None,
                    should_speak=True,
                    tone="neutral"
                ),
                follow_up_probe=None,
                transition=None
            ),
            next_question=None,
            flow_control=FlowControlModel(
                should_proceed_to_next=True,
                needs_follow_up=False,
                quality_sufficient=True,
                conversation_stage="mid"
            ),
            error=str(e)
        )


@app.post("/api/interview/submit-followup", response_model=SubmitFollowUpResponse)
async def submit_followup(request: SubmitFollowUpRequest):
    """
    Submit answer to follow-up probe.

    Used when AI asked for more detail and user elaborates.
    Combines with original answer for quality assessment.
    """
    logging.info(f"POST /api/interview/submit-followup - Session: {request.session_id}, Q{request.original_question_id}")

    try:
        # Validate inputs
        validate_uuid(request.session_id, "session_id")
        difficulty = validate_difficulty(request.difficulty)

        # Get orchestrator
        orchestrator = get_orchestrator()

        # Handle follow-up answer
        result = await orchestrator.handle_follow_up_answer(
            session_id=UUID(request.session_id),
            follow_up_answer=request.follow_up_answer,
            original_question_id=request.original_question_id,
            original_question_text=request.original_question_text,
            original_question_intent=request.original_question_intent,
            role=request.role,
            difficulty=difficulty,
            total_questions=request.total_questions,
            generate_audio=request.generate_audio
        )

        # Build response models
        ai_response_data = result.get("ai_response", {})
        ai_response = AIResponseModel(
            acknowledgment=AcknowledgmentModel(**ai_response_data["acknowledgment"]) if ai_response_data.get("acknowledgment") else None,
            follow_up_probe=FollowUpProbeModel(**ai_response_data["follow_up_probe"]) if ai_response_data.get("follow_up_probe") else None,
            transition=TransitionModel(**ai_response_data["transition"]) if ai_response_data.get("transition") else None
        )

        flow_control_data = result.get("flow_control", {})
        flow_control = FlowControlModel(
            should_proceed_to_next=flow_control_data.get("should_proceed_to_next", True),
            needs_follow_up=flow_control_data.get("needs_follow_up", False),
            quality_sufficient=flow_control_data.get("quality_sufficient", True),
            conversation_stage=flow_control_data.get("conversation_stage", "mid")
        )

        # Build next question model if present
        next_question = None
        if result.get("next_question"):
            nq = result["next_question"]
            q = nq.get("question", {})
            next_question = NextQuestionWithAudioModel(
                question=QuestionWithAudioModel(
                    id=q.get("id", request.original_question_id + 1),
                    text=q.get("text", ""),
                    intent=q.get("intent", "general"),
                    type=q.get("type", "standard"),
                    audio_url=q.get("audio_url")
                ),
                interviewer_comment=nq.get("interviewer_comment"),
                interviewer_comment_audio_url=nq.get("interviewer_comment_audio_url"),
                references=nq.get("references"),
                metadata=nq.get("metadata")
            )

        logging.info(f"Follow-up processed - Proceed: {flow_control.should_proceed_to_next}")

        return SubmitFollowUpResponse(
            follow_up_stored=True,
            answer_id=result.get("answer_id"),
            is_follow_up_response=True,
            original_question_id=request.original_question_id,
            follow_up_attempt=result.get("follow_up_attempt", 1),
            ai_response=ai_response,
            next_question=next_question,
            flow_control=flow_control,
            quality_metrics=result.get("quality_metrics"),
            error=result.get("error")
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error in submit-followup: {str(e)}")
        return SubmitFollowUpResponse(
            follow_up_stored=False,
            answer_id=None,
            is_follow_up_response=True,
            original_question_id=request.original_question_id,
            follow_up_attempt=1,
            ai_response=AIResponseModel(
                acknowledgment=AcknowledgmentModel(
                    text="Thank you for that clarification.",
                    audio_url=None,
                    should_speak=True,
                    tone="neutral"
                ),
                follow_up_probe=None,
                transition=None
            ),
            next_question=None,
            flow_control=FlowControlModel(
                should_proceed_to_next=True,
                needs_follow_up=False,
                quality_sufficient=True,
                conversation_stage="mid"
            ),
            error=str(e)
        )


@app.post("/api/interview/start-with-audio", response_model=StartInterviewWithAudioResponse)
async def start_interview_with_audio(request: StartInterviewWithAudioRequest):
    """
    Start interview and get first question with audio.

    Enhanced version of interview start with TTS support.
    Returns the first question with optional audio URL.
    """
    logging.info(f"POST /api/interview/start-with-audio - Session: {request.session_id}, Role: {request.role}")

    try:
        # Validate inputs
        validate_uuid(request.session_id, "session_id")
        difficulty = validate_difficulty(request.difficulty)

        # Get orchestrator
        orchestrator = get_orchestrator()

        # Generate first question with audio
        first_question_result = await orchestrator.generate_question_with_audio(
            session_id=UUID(request.session_id),
            current_question_number=1,
            role=request.role,
            difficulty=difficulty,
            total_questions=request.total_questions,
            generate_audio=request.generate_audio
        )

        # Build response
        q = first_question_result.get("question", {})
        first_question = NextQuestionWithAudioModel(
            question=QuestionWithAudioModel(
                id=q.get("id", 1),
                text=q.get("text", "Tell me about yourself."),
                intent=q.get("intent", "introduction"),
                type=q.get("type", "standard"),
                audio_url=q.get("audio_url")
            ),
            interviewer_comment=first_question_result.get("interviewer_comment"),
            interviewer_comment_audio_url=first_question_result.get("interviewer_comment_audio_url"),
            references=first_question_result.get("references"),
            metadata=first_question_result.get("metadata")
        )

        logging.info(f"Interview started with audio - Q1 audio: {q.get('audio_url')}")

        return StartInterviewWithAudioResponse(
            interview_started=True,
            session_id=request.session_id,
            first_question=first_question
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error in start-with-audio: {str(e)}")
        return StartInterviewWithAudioResponse(
            interview_started=False,
            session_id=request.session_id,
            first_question=None,
            error=str(e)
        )


@app.post("/api/interview/regenerate-audio", response_model=RegenerateAudioResponse)
async def regenerate_audio(request: RegenerateAudioRequest):
    """
    Regenerate audio for existing text.

    Useful for changing voices, fixing audio issues, or cache refresh.
    """
    logging.info(f"POST /api/interview/regenerate-audio - Context: {request.context_type}, Force: {request.force_regenerate}")

    try:
        tts = get_tts()

        # Check if should speak this text
        if not tts.should_speak_this(request.text, request.context_type):
            return RegenerateAudioResponse(
                audio_generated=False,
                audio_url=None,
                cached=False,
                error="Text should not be spoken (too long, code, or system message)"
            )

        # Generate audio
        kwargs = {}
        if request.voice:
            kwargs["voice"] = request.voice

        if request.force_regenerate:
            # Direct generation without cache
            audio_bytes = await tts.generate_for_interview_context(
                request.text,
                context_type=request.context_type,
                conversation_stage=request.conversation_stage
            )
            # Save with unique name
            import uuid
            filename = f"regen_{uuid.uuid4().hex[:8]}"
            file_path = await tts.save_audio_file(audio_bytes, filename)
            audio_url = f"/api/audio/{Path(file_path).name}"
            cached = False
        else:
            # Use cache
            cache_key = f"regen_{hash(request.text) % 100000}"
            audio_bytes, file_path = await tts.generate_and_cache(
                request.text,
                cache_key,
                **kwargs
            )
            audio_url = f"/api/audio/{Path(file_path).name}"
            cached = True

        logging.info(f"Audio regenerated: {audio_url}")

        return RegenerateAudioResponse(
            audio_generated=True,
            audio_url=audio_url,
            cached=cached
        )

    except Exception as e:
        logging.exception(f"Error regenerating audio: {str(e)}")
        return RegenerateAudioResponse(
            audio_generated=False,
            audio_url=None,
            cached=False,
            error=str(e)
        )


@app.get("/api/interview/audio-status/{session_id}", response_model=AudioStatusResponse)
async def get_audio_status(session_id: str):
    """
    Get audio generation status and cache statistics.

    Shows cache size and stats for monitoring.
    """
    logging.info(f"GET /api/interview/audio-status/{session_id}")

    try:
        validate_uuid(session_id, "session_id")

        tts = get_tts()
        cache_stats = tts.get_cache_stats()

        return AudioStatusResponse(
            session_id=session_id,
            cache_stats=cache_stats
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error getting audio status: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/audio/cache/clear", response_model=ClearCacheResponse)
async def clear_audio_cache(request: ClearCacheRequest):
    """
    Clear old audio cache files.

    Removes audio files older than specified days.
    Admin/maintenance endpoint.
    """
    logging.info(f"DELETE /api/audio/cache/clear - Older than {request.older_than_days} days")

    try:
        from datetime import datetime, timedelta
        import os

        cache_dir = Path("audio_cache")
        if not cache_dir.exists():
            return ClearCacheResponse(files_deleted=0, space_freed_mb=0.0)

        cutoff_time = datetime.now() - timedelta(days=request.older_than_days)
        files_deleted = 0
        space_freed = 0

        for file_path in cache_dir.glob("*.mp3"):
            try:
                file_stat = file_path.stat()
                file_mtime = datetime.fromtimestamp(file_stat.st_mtime)

                if file_mtime < cutoff_time:
                    space_freed += file_stat.st_size
                    file_path.unlink()
                    files_deleted += 1

            except Exception as e:
                logging.warning(f"Failed to delete {file_path}: {e}")

        space_freed_mb = round(space_freed / (1024 * 1024), 2)

        logging.info(f"Cache cleared: {files_deleted} files, {space_freed_mb} MB freed")

        return ClearCacheResponse(
            files_deleted=files_deleted,
            space_freed_mb=space_freed_mb
        )

    except Exception as e:
        logging.exception(f"Error clearing cache: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== Job Introduction Generator Endpoints ====================

class JobDescriptionInput(BaseModel):
    """Job description input for introduction generation."""
    company_name: str = Field(..., description="Name of the company")
    job_title: str = Field(..., description="Job title/position")
    team_name: Optional[str] = Field(None, description="Team name if applicable")
    location: Optional[str] = Field(None, description="Job location")
    responsibilities: List[str] = Field(default=[], description="List of job responsibilities")
    requirements: List[str] = Field(default=[], description="List of required qualifications")
    nice_to_have: List[str] = Field(default=[], description="List of preferred qualifications")
    company_description: Optional[str] = Field(None, description="Brief company description")
    role_description: Optional[str] = Field(None, description="Full role description text")


class GenerateIntroductionRequest(BaseModel):
    """Request for generating job introduction."""
    job_description: JobDescriptionInput
    candidate_name: Optional[str] = Field(None, description="Candidate name for personalization")
    candidate_resume_summary: Optional[str] = Field(None, description="Resume summary for personalization")
    mode: str = Field(default="concise", description="Introduction mode: 'concise' (30-45s) or 'detailed' (60-90s)")
    include_first_question: bool = Field(default=True, description="Include personalized first question")
    generate_audio: bool = Field(default=True, description="Generate TTS audio for segments")


class IntroductionSegment(BaseModel):
    """A single introduction segment."""
    segment_type: str
    text: str
    order: int
    duration_estimate_seconds: float
    audio_url: Optional[str] = None


class FirstQuestionOutput(BaseModel):
    """First question output."""
    segment_type: str
    text: str
    order: int
    duration_estimate_seconds: float
    audio_url: Optional[str] = None


class GenerateIntroductionResponse(BaseModel):
    """Response from introduction generation."""
    segments: List[IntroductionSegment]
    first_question: Optional[FirstQuestionOutput] = None
    total_duration_seconds: float
    mode: str
    error: Optional[str] = None


@app.post("/api/interview/generate-introduction", response_model=GenerateIntroductionResponse)
async def generate_interview_introduction(request: GenerateIntroductionRequest):
    """
    Generate a warm, professional introduction for a job-specific interview.

    This creates a complete opening sequence including:
    - Greeting (personalized if name provided)
    - Role overview
    - Key responsibilities
    - Requirements summary
    - Transition to interview

    Each segment can include TTS audio.

    Example request:
    ```json
    {
        "job_description": {
            "company_name": "Google",
            "job_title": "Senior Software Engineer",
            "team_name": "Backend Infrastructure Team",
            "responsibilities": ["Build scalable systems", "Design APIs"],
            "requirements": ["5+ years Python", "System design skills"]
        },
        "candidate_name": "Alex",
        "mode": "concise",
        "generate_audio": true
    }
    ```
    """
    logging.info(f"POST /api/interview/generate-introduction - {request.job_description.job_title} at {request.job_description.company_name}")

    try:
        # Convert Pydantic model to dict
        job_dict = {
            "company_name": request.job_description.company_name,
            "job_title": request.job_description.job_title,
            "team_name": request.job_description.team_name,
            "location": request.job_description.location,
            "responsibilities": request.job_description.responsibilities,
            "requirements": request.job_description.requirements,
            "nice_to_have": request.job_description.nice_to_have,
            "company_description": request.job_description.company_description,
            "role_description": request.job_description.role_description,
        }

        # Generate introduction
        result = await generate_job_introduction(
            job_description=job_dict,
            candidate_name=request.candidate_name,
            candidate_resume_summary=request.candidate_resume_summary,
            mode=request.mode,
            include_first_question=request.include_first_question,
            generate_audio=request.generate_audio
        )

        # Convert segments to response model
        segments = [
            IntroductionSegment(
                segment_type=s.get("segment_type", ""),
                text=s.get("text", ""),
                order=s.get("order", 0),
                duration_estimate_seconds=s.get("duration_estimate_seconds", 0),
                audio_url=s.get("audio_url")
            )
            for s in result.get("segments", [])
        ]

        # Convert first question if present
        first_question = None
        if result.get("first_question"):
            fq = result["first_question"]
            first_question = FirstQuestionOutput(
                segment_type=fq.get("segment_type", "first_question"),
                text=fq.get("text", ""),
                order=fq.get("order", 0),
                duration_estimate_seconds=fq.get("duration_estimate_seconds", 0),
                audio_url=fq.get("audio_url")
            )

        return GenerateIntroductionResponse(
            segments=segments,
            first_question=first_question,
            total_duration_seconds=result.get("total_duration_seconds", 0),
            mode=result.get("mode", request.mode)
        )

    except Exception as e:
        logging.exception(f"Error generating introduction: {str(e)}")
        return GenerateIntroductionResponse(
            segments=[],
            first_question=None,
            total_duration_seconds=0,
            mode=request.mode,
            error=str(e)
        )


class GenerateFirstQuestionRequest(BaseModel):
    """Request for generating personalized first question only."""
    job_description: JobDescriptionInput
    candidate_resume_summary: Optional[str] = Field(None, description="Resume summary for personalization")
    generate_audio: bool = Field(default=True, description="Generate TTS audio")


class GenerateFirstQuestionResponse(BaseModel):
    """Response with personalized first question."""
    question_text: str
    audio_url: Optional[str] = None
    duration_estimate_seconds: float = 8.0
    error: Optional[str] = None


@app.post("/api/interview/generate-first-question", response_model=GenerateFirstQuestionResponse)
async def generate_personalized_first_question(request: GenerateFirstQuestionRequest):
    """
    Generate a personalized first interview question tied to the job description.

    Instead of generic "Tell me about yourself", creates a question
    that immediately ties to the role.

    Example output:
    "Tell me about yourself and specifically what attracted you to
    the Senior Software Engineer position at Google."
    """
    logging.info(f"POST /api/interview/generate-first-question - {request.job_description.job_title}")

    try:
        # Convert to dict
        job_dict = {
            "company_name": request.job_description.company_name,
            "job_title": request.job_description.job_title,
            "requirements": request.job_description.requirements,
        }

        # Get generator
        generator = get_job_introduction_generator()

        # Generate question
        question_text = await generator.generate_personalized_first_question(
            job_description=job_dict,
            candidate_resume_summary=request.candidate_resume_summary
        )

        audio_url = None

        # Generate audio if requested
        if request.generate_audio:
            try:
                tts = get_tts_service()
                import hashlib
                text_hash = hashlib.md5(question_text.encode()).hexdigest()[:8]
                cache_key = f"first_question_{text_hash}"

                audio_bytes, audio_path = await tts.generate_and_cache(
                    text=question_text,
                    cache_key=cache_key,
                    voice="alloy",
                    speed=0.9
                )

                if audio_path:
                    import os
                    filename = os.path.basename(audio_path)
                    audio_url = f"/api/audio/{filename}"

            except Exception as e:
                logging.warning(f"Failed to generate audio for first question: {e}")

        return GenerateFirstQuestionResponse(
            question_text=question_text,
            audio_url=audio_url,
            duration_estimate_seconds=8.0
        )

    except Exception as e:
        logging.exception(f"Error generating first question: {str(e)}")
        return GenerateFirstQuestionResponse(
            question_text="Tell me about yourself.",
            audio_url=None,
            error=str(e)
        )


# ==================== Job Description Interview Start Endpoints ====================

# Dependency injection helpers
def get_orchestrator() -> InterviewOrchestrator:
    """Get or create InterviewOrchestrator instance."""
    return InterviewOrchestrator()


def get_job_intro_generator() -> JobIntroductionGenerator:
    """Get or create JobIntroductionGenerator instance."""
    return get_job_introduction_generator()


# Pydantic models for job description interview
class JobDescriptionData(BaseModel):
    """Job description data for starting an interview."""
    company_name: Optional[str] = Field(None, description="Name of the company")
    company_description: Optional[str] = Field(None, description="Brief company description")
    job_title: str = Field(..., description="Job title/position")
    team_name: Optional[str] = Field(None, description="Team name if applicable")
    location: Optional[str] = Field(None, description="Job location")
    responsibilities: List[str] = Field(default=[], description="List of job responsibilities")
    requirements: List[str] = Field(default=[], description="List of required qualifications")
    nice_to_have: List[str] = Field(default=[], description="List of preferred qualifications")
    role_description: Optional[str] = Field(None, description="Full job description text")


class StartInterviewWithJDRequest(BaseModel):
    """Request to start an interview with job description."""
    role: str = Field(..., description="Role being interviewed for")
    difficulty: str = Field(default="medium", description="Interview difficulty level")
    job_description: JobDescriptionData = Field(..., description="Job description details")
    candidate_name: Optional[str] = Field(None, description="Candidate name for personalization")
    candidate_resume_summary: Optional[str] = Field(None, description="Resume summary for personalization")
    generate_audio: bool = Field(default=True, description="Generate TTS audio for introduction")
    introduction_mode: str = Field(default="concise", description="Introduction mode: 'concise' or 'detailed'")
    question_count: int = Field(default=10, description="Total number of interview questions")


class IntroSegmentOutput(BaseModel):
    """Output for a single introduction segment."""
    segment_type: str
    text: str
    order: int
    duration_estimate_seconds: float
    audio_url: Optional[str] = None


class FirstQuestionData(BaseModel):
    """First question data with audio."""
    question_id: int = 1
    text: str
    intent: str = "introduction"
    type: str = "standard"
    audio_url: Optional[str] = None
    duration_estimate_seconds: float = 8.0


class JobDescriptionStoredOutput(BaseModel):
    """Stored job description output."""
    id: str
    company_name: Optional[str] = None
    job_title: str
    created_at: str


class StartInterviewWithJDResponse(BaseModel):
    """Response from starting interview with job description."""
    session_id: str
    job_description_id: str
    introduction_sequence: List[IntroSegmentOutput]
    first_question: Optional[FirstQuestionData] = None
    estimated_intro_duration_seconds: float
    job_description_stored: JobDescriptionStoredOutput
    error: Optional[str] = None


def store_job_description_to_db(session_id: str, jd_data: JobDescriptionData) -> JobDescription:
    """
    Store job description data to database.

    Args:
        session_id: Interview session ID
        jd_data: Job description data

    Returns:
        Created JobDescription record
    """
    with Session(engine) as db_session:
        job_desc = JobDescription(
            session_id=session_id,
            company_name=jd_data.company_name,
            company_description=jd_data.company_description,
            job_title=jd_data.job_title,
            team_name=jd_data.team_name,
            location=jd_data.location,
            responsibilities=json.dumps(jd_data.responsibilities) if jd_data.responsibilities else None,
            requirements=json.dumps(jd_data.requirements) if jd_data.requirements else None,
            nice_to_have=json.dumps(jd_data.nice_to_have) if jd_data.nice_to_have else None,
            role_description=jd_data.role_description
        )

        db_session.add(job_desc)
        db_session.commit()
        db_session.refresh(job_desc)

        return job_desc


def create_interview_session_for_jd(
    role: str,
    difficulty: str,
    question_count: int,
    job_title: str
) -> InterviewSession:
    """
    Create a new interview session for job description interview.

    Args:
        role: Role being interviewed for
        difficulty: Interview difficulty
        question_count: Number of questions
        job_title: Job title from JD

    Returns:
        Created InterviewSession record
    """
    with Session(engine) as db_session:
        session = InterviewSession(
            role=role,
            difficulty=difficulty,
            question_count=question_count,
            interviewer_names=["AI Interviewer"],
            plan={"job_title": job_title, "type": "job_description_interview"}
        )

        db_session.add(session)
        db_session.commit()
        db_session.refresh(session)

        return session


@app.post("/api/interview/start-with-job-description", response_model=StartInterviewWithJDResponse)
async def start_interview_with_job_description(
    request: StartInterviewWithJDRequest,
    orchestrator: InterviewOrchestrator = Depends(get_orchestrator),
    job_intro_gen: JobIntroductionGenerator = Depends(get_job_intro_generator)
):
    """
    Start an interview with a job description introduction.

    This endpoint:
    1. Creates an interview session
    2. Stores the job description
    3. Generates a warm introduction sequence with audio
    4. Generates a personalized first question

    The introduction sequence includes:
    - Greeting (personalized if name provided)
    - Role overview
    - Responsibilities summary
    - Requirements summary
    - Transition to interview

    Example request:
    ```json
    {
        "role": "Software Engineer",
        "difficulty": "medium",
        "job_description": {
            "company_name": "Google",
            "job_title": "Senior Software Engineer",
            "team_name": "Backend Infrastructure",
            "responsibilities": ["Build scalable systems", "Design APIs"],
            "requirements": ["5+ years Python", "System design skills"]
        },
        "candidate_name": "Alex",
        "generate_audio": true
    }
    ```
    """
    logging.info(f"POST /api/interview/start-with-job-description - {request.job_description.job_title} at {request.job_description.company_name or 'Unknown Company'}")

    try:
        # 1. Create interview session
        session = create_interview_session_for_jd(
            role=request.role,
            difficulty=request.difficulty,
            question_count=request.question_count,
            job_title=request.job_description.job_title
        )
        logging.info(f"Created interview session: {session.id}")

        # 2. Store job description
        job_desc = store_job_description_to_db(
            session_id=session.id,
            jd_data=request.job_description
        )
        logging.info(f"Stored job description: {job_desc.id}")

        # 3. Convert JD to dict for generators
        jd_dict = {
            "company_name": request.job_description.company_name,
            "company_description": request.job_description.company_description,
            "job_title": request.job_description.job_title,
            "team_name": request.job_description.team_name,
            "location": request.job_description.location,
            "responsibilities": request.job_description.responsibilities,
            "requirements": request.job_description.requirements,
            "nice_to_have": request.job_description.nice_to_have,
            "role_description": request.job_description.role_description,
        }

        # 4. Generate introduction sequence
        introduction_segments = await job_intro_gen.generate_opening_sequence(
            job_description=jd_dict,
            candidate_name=request.candidate_name,
            candidate_resume_summary=request.candidate_resume_summary,
            generate_audio=request.generate_audio
        )
        logging.info(f"Generated {len(introduction_segments)} introduction segments")

        # 5. Generate personalized first question
        first_question_text = await job_intro_gen.generate_personalized_first_question(
            job_description=jd_dict,
            candidate_resume_summary=request.candidate_resume_summary
        )
        logging.info(f"Generated first question: {first_question_text[:50]}...")

        # 6. Generate audio for first question if requested
        first_question_audio_url = None
        if request.generate_audio:
            try:
                tts = get_tts_service()
                import hashlib
                text_hash = hashlib.md5(first_question_text.encode()).hexdigest()[:8]
                cache_key = f"jd_first_question_{session.id[:8]}_{text_hash}"

                audio_bytes, audio_path = await tts.generate_and_cache(
                    text=first_question_text,
                    cache_key=cache_key,
                    voice="alloy",
                    speed=0.9
                )

                if audio_path:
                    filename = os.path.basename(audio_path)
                    first_question_audio_url = f"/api/audio/{filename}"

            except Exception as e:
                logging.warning(f"Failed to generate audio for first question: {e}")

        # 7. Build response
        intro_output = [
            IntroSegmentOutput(
                segment_type=s.get("segment_type", ""),
                text=s.get("text", ""),
                order=s.get("order", 0),
                duration_estimate_seconds=s.get("duration_estimate_seconds", 5),
                audio_url=s.get("audio_url")
            )
            for s in introduction_segments
        ]

        estimated_duration = sum(s.get("duration_estimate_seconds", 5) for s in introduction_segments)

        first_question_output = FirstQuestionData(
            question_id=1,
            text=first_question_text,
            intent="introduction",
            type="standard",
            audio_url=first_question_audio_url,
            duration_estimate_seconds=8.0
        )

        job_desc_output = JobDescriptionStoredOutput(
            id=job_desc.id,
            company_name=job_desc.company_name,
            job_title=job_desc.job_title,
            created_at=job_desc.created_at.isoformat()
        )

        return StartInterviewWithJDResponse(
            session_id=session.id,
            job_description_id=job_desc.id,
            introduction_sequence=intro_output,
            first_question=first_question_output,
            estimated_intro_duration_seconds=estimated_duration,
            job_description_stored=job_desc_output
        )

    except Exception as e:
        logging.exception(f"Error starting interview with job description: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to start interview: {str(e)}"
        )


# ==================== Job Description Parsing Endpoint ====================

class ParseJDRequest(BaseModel):
    """Request to parse a job description text."""
    raw_text: str = Field(..., description="Raw job description text to parse")
    source_type: str = Field(default="text", description="Source type: 'text', 'pdf', 'url'")


class ParsedJobDescription(BaseModel):
    """Parsed job description with extracted fields."""
    company_name: Optional[str] = None
    company_description: Optional[str] = None
    job_title: Optional[str] = None
    team_name: Optional[str] = None
    location: Optional[str] = None
    responsibilities: List[str] = []
    requirements: List[str] = []
    nice_to_have: List[str] = []
    salary_range: Optional[str] = None
    employment_type: Optional[str] = None
    experience_level: Optional[str] = None
    role_description: Optional[str] = None
    confidence_score: float = 0.0


class ParseJDResponse(BaseModel):
    """Response from job description parsing."""
    parsed: ParsedJobDescription
    raw_text_length: int
    parsing_method: str = "llm"
    error: Optional[str] = None


async def parse_job_description_with_llm(raw_text: str) -> Dict:
    """
    Parse job description text using LLM to extract structured data.

    Args:
        raw_text: Raw job description text

    Returns:
        Dict with parsed fields
    """
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    prompt = f"""Parse this job description and extract structured information.

JOB DESCRIPTION:
{raw_text[:4000]}  # Limit to 4000 chars to stay within token limits

Extract the following fields and return as JSON:
{{
    "company_name": "Company name if mentioned",
    "company_description": "Brief company description if available",
    "job_title": "The job title/position",
    "team_name": "Team or department name if mentioned",
    "location": "Job location (city, remote, hybrid, etc.)",
    "responsibilities": ["List", "of", "key", "responsibilities"],
    "requirements": ["List", "of", "required", "qualifications"],
    "nice_to_have": ["List", "of", "preferred/bonus", "qualifications"],
    "salary_range": "Salary range if mentioned",
    "employment_type": "Full-time, Part-time, Contract, etc.",
    "experience_level": "Entry, Mid, Senior, Lead, etc.",
    "confidence_score": 0.0-1.0 based on how well you could extract info
}}

Rules:
- Return ONLY valid JSON, no markdown or explanation
- Use null for fields you can't find
- Keep responsibilities and requirements concise (max 10 items each)
- Combine similar items
- Confidence score: 1.0 = all fields found clearly, 0.5 = partial, 0.0 = couldn't parse"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.3,
            max_tokens=1500,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert at parsing job descriptions. Extract structured data accurately. Return only valid JSON."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        result_text = response.choices[0].message.content.strip()

        # Handle potential markdown code blocks
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]

        parsed = json.loads(result_text)
        return parsed

    except json.JSONDecodeError as e:
        logging.error(f"Failed to parse LLM response as JSON: {e}")
        return {
            "job_title": "Unknown Position",
            "responsibilities": [],
            "requirements": [],
            "nice_to_have": [],
            "confidence_score": 0.0
        }
    except Exception as e:
        logging.exception(f"LLM parsing failed: {e}")
        raise


@app.post("/api/job-description/parse", response_model=ParseJDResponse)
async def parse_job_description(request: ParseJDRequest):
    """
    Parse uploaded job description text and extract structured data.

    This uses LLM to intelligently extract:
    - Company information
    - Job title and team
    - Location
    - Responsibilities
    - Requirements (required vs nice-to-have)
    - Salary and employment type

    Example request:
    ```json
    {
        "raw_text": "Senior Software Engineer at Google\\n\\nAbout the role...\\n\\nResponsibilities:\\n- Build scalable systems...",
        "source_type": "text"
    }
    ```
    """
    logging.info(f"POST /api/job-description/parse - {len(request.raw_text)} chars")

    if not request.raw_text or len(request.raw_text.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Job description text is too short. Please provide at least 50 characters."
        )

    try:
        # Parse using LLM
        parsed_data = await parse_job_description_with_llm(request.raw_text)

        # Build response
        parsed = ParsedJobDescription(
            company_name=parsed_data.get("company_name"),
            company_description=parsed_data.get("company_description"),
            job_title=parsed_data.get("job_title"),
            team_name=parsed_data.get("team_name"),
            location=parsed_data.get("location"),
            responsibilities=parsed_data.get("responsibilities", []) or [],
            requirements=parsed_data.get("requirements", []) or [],
            nice_to_have=parsed_data.get("nice_to_have", []) or [],
            salary_range=parsed_data.get("salary_range"),
            employment_type=parsed_data.get("employment_type"),
            experience_level=parsed_data.get("experience_level"),
            role_description=request.raw_text[:2000],  # Store first 2000 chars
            confidence_score=parsed_data.get("confidence_score", 0.5)
        )

        return ParseJDResponse(
            parsed=parsed,
            raw_text_length=len(request.raw_text),
            parsing_method="llm"
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error parsing job description: {str(e)}")
        return ParseJDResponse(
            parsed=ParsedJobDescription(
                job_title="Unknown Position",
                confidence_score=0.0
            ),
            raw_text_length=len(request.raw_text),
            parsing_method="llm",
            error=str(e)
        )


@app.post("/api/job-description/parse-file")
async def parse_job_description_file(file: UploadFile = File(...)):
    """
    Parse an uploaded job description file (PDF or text).

    Extracts text from the file and then parses it to extract
    structured job description data.

    Supported formats:
    - PDF (.pdf)
    - Text (.txt)
    - Word (.docx)
    """
    logging.info(f"POST /api/job-description/parse-file - {file.filename}")

    # Check file type
    filename = file.filename.lower() if file.filename else ""

    if not (filename.endswith('.pdf') or filename.endswith('.txt') or filename.endswith('.docx')):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF, TXT, or DOCX file."
        )

    try:
        # Read file content
        content = await file.read()

        # Extract text based on file type
        raw_text = ""

        if filename.endswith('.pdf'):
            # Extract text from PDF
            pdf_doc = fitz.open(stream=content, filetype="pdf")
            for page in pdf_doc:
                raw_text += page.get_text()
            pdf_doc.close()

        elif filename.endswith('.docx'):
            # Extract text from DOCX
            doc = docx.Document(io.BytesIO(content))
            raw_text = "\n".join(para.text for para in doc.paragraphs)

        else:  # .txt
            raw_text = content.decode('utf-8', errors='ignore')

        if len(raw_text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Could not extract sufficient text from the file. Please check the file content."
            )

        # Parse the extracted text
        parsed_data = await parse_job_description_with_llm(raw_text)

        # Build response
        parsed = ParsedJobDescription(
            company_name=parsed_data.get("company_name"),
            company_description=parsed_data.get("company_description"),
            job_title=parsed_data.get("job_title"),
            team_name=parsed_data.get("team_name"),
            location=parsed_data.get("location"),
            responsibilities=parsed_data.get("responsibilities", []) or [],
            requirements=parsed_data.get("requirements", []) or [],
            nice_to_have=parsed_data.get("nice_to_have", []) or [],
            salary_range=parsed_data.get("salary_range"),
            employment_type=parsed_data.get("employment_type"),
            experience_level=parsed_data.get("experience_level"),
            role_description=raw_text[:2000],
            confidence_score=parsed_data.get("confidence_score", 0.5)
        )

        return ParseJDResponse(
            parsed=parsed,
            raw_text_length=len(raw_text),
            parsing_method="llm"
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error parsing job description file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse file: {str(e)}"
        )


@app.get("/api/job-description/{session_id}")
async def get_job_description(session_id: str):
    """
    Get the stored job description for a session.

    Returns the full job description data associated with an interview session.
    """
    logging.info(f"GET /api/job-description/{session_id}")

    try:
        validate_uuid(session_id, "session_id")

        with Session(engine) as db_session:
            job_desc = db_session.exec(
                select(JobDescription).where(JobDescription.session_id == session_id)
            ).first()

            if not job_desc:
                raise HTTPException(
                    status_code=404,
                    detail=f"No job description found for session {session_id}"
                )

            return {
                "id": job_desc.id,
                "session_id": job_desc.session_id,
                "company_name": job_desc.company_name,
                "company_description": job_desc.company_description,
                "job_title": job_desc.job_title,
                "team_name": job_desc.team_name,
                "location": job_desc.location,
                "responsibilities": json.loads(job_desc.responsibilities) if job_desc.responsibilities else [],
                "requirements": json.loads(job_desc.requirements) if job_desc.requirements else [],
                "nice_to_have": json.loads(job_desc.nice_to_have) if job_desc.nice_to_have else [],
                "role_description": job_desc.role_description,
                "created_at": job_desc.created_at.isoformat()
            }

    except HTTPException:
        raise
    except Exception as e:
        logging.exception(f"Error getting job description: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
